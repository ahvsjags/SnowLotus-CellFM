from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
RELEASE = ROOT / "release_metadata"

REPO_URL = "https://github.com/ahvsjags/SnowLotus-CellFM"
BRANCH = "agent/remote-pipeline-20260728"
RELEASE_TAG = "v0.9.0-plant-general-lora"


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def fmt(value: float) -> str:
    return f"{value:.4f}"


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def curve_at(curve: list[dict[str, Any]], rate: float) -> dict[str, Any]:
    for item in curve:
        if abs(float(item.get("acceptance_rate", -1)) - rate) < 1e-9:
            return item
    return {}


def build_context() -> dict[str, Any]:
    comparison = read_json(RELEASE / "v9_benchmarks" / "v9_lora_vs_v3_shared_comparison.json")
    ontology = read_json(RELEASE / "species_ontology_label_benchmark_v9.json")
    biology = read_json(RELEASE / "plant_biology_case_study_v9.json")
    external = read_json(RELEASE / "external_benchmark_panel_v9.json")
    open_set = read_json(RELEASE / "open_set_calibration_v9.json")
    multi_case = read_json(RELEASE / "multispecies_scplantdb_case_v10.json")
    candidate = comparison["candidate"]["summary"]
    baseline = comparison["baseline"]["summary"]
    delta = comparison["delta"]
    ontology_action = ontology["protocols"]["leave_species_out_ontology_actionable"]
    marker_overview = biology["marker_overview"]
    adapter_layer = biology["adapter_layer"]
    return {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M Asia/Shanghai"),
        "candidate": candidate,
        "baseline": baseline,
        "delta": delta,
        "ontology_action": ontology_action,
        "ontology_label_summary": ontology["label_summary"],
        "marker_overview": marker_overview,
        "adapter_count": adapter_layer["adapter_count"],
        "external_summary": external["summary"],
        "api_top30": curve_at(open_set["api_head_confidence"]["fine_confidence_curve"], 0.3),
        "api_top40": curve_at(open_set["api_head_confidence"]["fine_confidence_curve"], 0.4),
        "multi_case": multi_case,
    }


def abstract(ctx: dict[str, Any]) -> str:
    candidate = ctx["candidate"]
    baseline = ctx["baseline"]
    ontology = ctx["ontology_action"]
    markers = ctx["marker_overview"]
    api_top30 = ctx["api_top30"]
    api_top40 = ctx["api_top40"]
    multi = ctx["multi_case"]
    return (
        "Plant single-cell and single-nucleus transcriptomic studies increasingly cover diverse species, "
        "tissues and assay formats, yet cross-study reuse is limited by heterogeneous matrix formats, "
        "non-unified cell-state names and species-specific gene identifiers. We present Plant-CellFM v9, "
        "a reproducible plant expression foundation-model and all-plant adapter framework for audited "
        "single-cell annotation. The release combines a public plant expression corpus, shared-gene "
        "Transformer representations, LoRA-based model freezing, runtime species-adapter resolution, "
        "hierarchical annotation outputs and server-side release verification. On the same shared-gene "
        "benchmark, Plant-CellFM v9 improves over the frozen v3 extended baseline in leave-dataset-out "
        f"all-cell accuracy ({fmt(candidate['leave_dataset_out']['fine']['accuracy_all'])} versus "
        f"{fmt(baseline['leave_dataset_out']['fine']['accuracy_all'])}) and leave-sample-out all-cell "
        f"accuracy ({fmt(candidate['leave_sample_out']['fine']['accuracy_all'])} versus "
        f"{fmt(baseline['leave_sample_out']['fine']['accuracy_all'])}). Under normalized leave-species-out "
        f"evaluation, v9 reaches all-cell accuracy {fmt(candidate['leave_species_out']['fine']['accuracy_all'])}, "
        f"coverage {fmt(candidate['leave_species_out']['fine']['coverage'])} and known-label accuracy "
        f"{fmt(candidate['leave_species_out']['fine']['accuracy'])}, supporting open-set cross-species "
        "transfer analysis rather than a universal high-accuracy claim. A plant cell-state ontology diagnostic "
        f"covers {ontology['n_test']:,} of {ontology['n_test_total']:,} cells "
        f"({pct(ontology['coverage'])}) after excluding unknown or unannotated states. The API confidence layer "
        f"reaches {pct(api_top30['selective_accuracy'])} and {pct(api_top40['selective_accuracy'])} selective "
        "accuracy when accepting the top 30% and 40% confidence cells. The release further "
        f"includes {ctx['adapter_count']} adapter entries, an Arabidopsis root case with "
        f"{markers['n_marker_rows']} marker-candidate rows across {markers['n_labels']} cell states and "
        f"{markers['root_identity_label_count']} root-identity states, and a multi-species scPlantDB case with "
        f"{multi['corpus']['cells']:,} cells across {multi['corpus']['species']} species. Plant-CellFM v9 therefore provides "
        "a traceable method and resource for plant single-cell annotation, benchmark auditing and target-species "
        "adapter transfer."
    )


def synopsis_markdown(ctx: dict[str, Any]) -> str:
    candidate = ctx["candidate"]
    baseline = ctx["baseline"]
    ontology = ctx["ontology_action"]
    markers = ctx["marker_overview"]
    external = ctx["external_summary"]
    api_top30 = ctx["api_top30"]
    api_top40 = ctx["api_top40"]
    multi = ctx["multi_case"]
    lines = [
        "# Plant-CellFM v9 English Submission Synopsis",
        "",
        f"Generated: `{ctx['generated_at']}`",
        "",
        f"Repository: {REPO_URL}",
        "",
        f"Branch: `{BRANCH}`",
        "",
        f"Release tag: `{RELEASE_TAG}`",
        "",
        "## Proposed Title",
        "",
        "Plant-CellFM: a reproducible foundation-model and adapter framework for plant single-cell annotation",
        "",
        "## Abstract",
        "",
        abstract(ctx),
        "",
        "## Significance Statement",
        "",
        "Plant single-cell atlases are expanding faster than their annotation conventions can be harmonized. "
        "Plant-CellFM v9 turns this practical bottleneck into a reproducible modelling problem: matrices, "
        "labels, adapters, checkpoints, benchmark splits and server health are all exposed as auditable release "
        "objects. Its central contribution is a reusable plant-general framework that makes cross-dataset transfer, "
        "open-set species transfer and target-species adapter preparation inspectable from the same code path.",
        "",
        "## Highlights",
        "",
        "- Plant-general foundation model for single-cell and single-nucleus plant expression annotation.",
        f"- All-plant adapter framework with {ctx['adapter_count']} adapter entries and universal fallback resolution.",
        "- Strict grouped evaluation, including leave-dataset-out, leave-sample-out and normalized leave-species-out protocols.",
        f"- v9 improves over frozen v3 in leave-dataset-out all-cell accuracy ({fmt(candidate['leave_dataset_out']['fine']['accuracy_all'])} versus {fmt(baseline['leave_dataset_out']['fine']['accuracy_all'])}) and leave-sample-out all-cell accuracy ({fmt(candidate['leave_sample_out']['fine']['accuracy_all'])} versus {fmt(baseline['leave_sample_out']['fine']['accuracy_all'])}).",
        f"- Ontology-actionable benchmark separates {pct(ontology['coverage'])} covered cells from unknown or unannotated states.",
        f"- Open-set calibration reaches {pct(api_top30['selective_accuracy'])}/{pct(api_top40['selective_accuracy'])} selective accuracy at top-30/top-40 confidence acceptance.",
        f"- Arabidopsis root case provides {markers['n_marker_rows']} marker-candidate rows across {markers['n_labels']} cell states.",
        f"- Multi-species scPlantDB case adds {multi['corpus']['cells']:,} cells across {multi['corpus']['species']} species and {multi['marker_record_count']} marker-candidate records.",
        "",
        "## Graphical Abstract Text",
        "",
        "Panel 1: Heterogeneous public plant matrices enter an audited corpus layer with accession, species, label and file-integrity records.",
        "",
        "Panel 2: Shared-gene expression profiles are encoded by the Plant-CellFM representation model and frozen through a LoRA release checkpoint.",
        "",
        "Panel 3: Runtime adapter resolution selects exact species adapters when available and falls back to a plant-universal adapter for new species.",
        "",
        "Panel 4: Grouped benchmarks quantify leave-dataset, leave-sample and open-set leave-species transfer against frozen v3, centroid and Seurat comparators.",
        "",
        "Panel 5: The Arabidopsis root case links model output to cell-state labels and marker-candidate mining for downstream biological interpretation.",
        "",
        "Panel 6: Open-set confidence calibration and the multi-species scPlantDB case show how high-confidence predictions, review routing and public-data biology examples are packaged for reuse.",
        "",
        "## Evidence At A Glance",
        "",
        f"- Completed metric rows in external benchmark panel: {external['completed_metric_rows']} / {external['rows']}.",
        f"- Completed formal comparisons in the current package: {external['completed_formal_comparisons']}.",
        f"- Normalized leave-species-out all-cell accuracy: {fmt(candidate['leave_species_out']['fine']['accuracy_all'])}.",
        f"- Normalized leave-species-out known-label accuracy: {fmt(candidate['leave_species_out']['fine']['accuracy'])}.",
        f"- Ontology-label actionable all-cell accuracy: {pct(ontology['accuracy_all'])}.",
        f"- Ontology-label known-label accuracy: {pct(ontology['accuracy'])}.",
        f"- Ontology-label macro-F1: {fmt(ontology['macro_f1'])}.",
        f"- API confidence top-30 selective accuracy: {pct(api_top30['selective_accuracy'])}.",
        f"- API confidence top-40 selective accuracy: {pct(api_top40['selective_accuracy'])}.",
        f"- Multi-species scPlantDB case: {multi['corpus']['cells']:,} cells, {multi['corpus']['species']} species, {multi['marker_record_count']} marker-candidate records.",
        "",
        "## Editorial Positioning",
        "",
        "The manuscript is positioned as a computational method/resource paper for plant single-cell annotation. "
        "The core promise is reproducibility, adapter-based plant generalization and transparent benchmark auditing. "
        "The submission reports open-set leave-species performance as diagnostic transfer evidence, adds selective "
        "annotation evidence for high-confidence predictions, treats Snow Lotus as one target-species adapter entry "
        "point and records scPlantLLM/scPlantAnnotate through official-source benchmark contracts pending official "
        "metric closure.",
        "",
        "## Submission Checklist",
        "",
        "- Use `SUBMISSION_INDEX_v9.md` as the reviewer entry point.",
        "- Use `manuscript/Plant_CellFM_v9_final_submission_zh_v1.docx` as the current full Chinese manuscript.",
        "- Use `manuscript/Plant_CellFM_v9_cover_letter.docx` as the cover letter.",
        "- Use this synopsis for the English abstract, highlights, significance and graphical abstract text.",
        "- Use `release_metadata/data_code_availability_v9.md` for repository, release and server reproducibility statements.",
        "- Verify the final editor package with `scripts/verify_v9_server_release.py` before resubmission.",
        "",
    ]
    return "\n".join(lines)


def highlights_json(ctx: dict[str, Any]) -> dict[str, Any]:
    candidate = ctx["candidate"]
    baseline = ctx["baseline"]
    ontology = ctx["ontology_action"]
    markers = ctx["marker_overview"]
    return {
        "schema_version": "plant_cellfm_v9_submission_highlights_v1",
        "generated_at": ctx["generated_at"],
        "repository": REPO_URL,
        "branch": BRANCH,
        "release_tag": RELEASE_TAG,
        "proposed_title": "Plant-CellFM: a reproducible foundation-model and adapter framework for plant single-cell annotation",
        "highlights": [
            "Plant-general foundation model for single-cell and single-nucleus plant expression annotation.",
            f"All-plant adapter framework with {ctx['adapter_count']} adapter entries and universal fallback resolution.",
            "Strict grouped evaluation separates leave-dataset, leave-sample and open-set leave-species transfer.",
            "Frozen v9 improves over frozen v3 on the shared-gene benchmark in leave-dataset-out and leave-sample-out protocols.",
            "Plant cell-state ontology diagnostics separate actionable labels from unknown or unannotated states.",
            "Open-set calibration provides a confidence-aware accept/review protocol for high-confidence annotations.",
            "Arabidopsis root case links adapter resolution, hierarchical annotation and marker-candidate mining.",
            "Multi-species scPlantDB case broadens the public-data biology demonstration beyond Arabidopsis.",
        ],
        "headline_numbers": {
            "leave_dataset_out_v9_accuracy_all": candidate["leave_dataset_out"]["fine"]["accuracy_all"],
            "leave_dataset_out_v3_accuracy_all": baseline["leave_dataset_out"]["fine"]["accuracy_all"],
            "leave_sample_out_v9_accuracy_all": candidate["leave_sample_out"]["fine"]["accuracy_all"],
            "leave_sample_out_v3_accuracy_all": baseline["leave_sample_out"]["fine"]["accuracy_all"],
            "leave_species_out_v9_accuracy_all": candidate["leave_species_out"]["fine"]["accuracy_all"],
            "leave_species_out_v9_coverage": candidate["leave_species_out"]["fine"]["coverage"],
            "leave_species_out_v9_known_label_accuracy": candidate["leave_species_out"]["fine"]["accuracy"],
            "ontology_actionable_coverage": ontology["coverage"],
            "ontology_actionable_accuracy_all": ontology["accuracy_all"],
            "ontology_known_label_accuracy": ontology["accuracy"],
            "ontology_macro_f1": ontology["macro_f1"],
            "api_confidence_top30_selective_accuracy": ctx["api_top30"]["selective_accuracy"],
            "api_confidence_top40_selective_accuracy": ctx["api_top40"]["selective_accuracy"],
            "adapter_count": ctx["adapter_count"],
            "arabidopsis_marker_candidate_rows": markers["n_marker_rows"],
            "arabidopsis_cell_states": markers["n_labels"],
            "arabidopsis_root_identity_states": markers["root_identity_label_count"],
            "multispecies_scplantdb_cells": ctx["multi_case"]["corpus"]["cells"],
            "multispecies_scplantdb_species": ctx["multi_case"]["corpus"]["species"],
            "multispecies_scplantdb_marker_candidates": ctx["multi_case"]["marker_record_count"],
        },
        "claim_safe_position": (
            "Use Plant-CellFM v9 as a plant-general reproducible method and resource. "
            "State leave-species-out performance as open-set transfer evidence, not universal high-accuracy annotation."
        ),
    }


def highlights_markdown(data: dict[str, Any]) -> str:
    numbers = data["headline_numbers"]
    lines = [
        "# Plant-CellFM v9 Submission Highlights",
        "",
        f"Generated: `{data['generated_at']}`",
        "",
        "## Proposed Title",
        "",
        data["proposed_title"],
        "",
        "## Highlights",
        "",
    ]
    for item in data["highlights"]:
        lines.append(f"- {item}")
    lines.extend(
        [
            "",
            "## Headline Numbers",
            "",
            f"- Leave-dataset-out all-cell accuracy: v9 {fmt(numbers['leave_dataset_out_v9_accuracy_all'])}; v3 {fmt(numbers['leave_dataset_out_v3_accuracy_all'])}.",
            f"- Leave-sample-out all-cell accuracy: v9 {fmt(numbers['leave_sample_out_v9_accuracy_all'])}; v3 {fmt(numbers['leave_sample_out_v3_accuracy_all'])}.",
            f"- Normalized leave-species-out all-cell accuracy: {fmt(numbers['leave_species_out_v9_accuracy_all'])}.",
            f"- Normalized leave-species-out coverage: {fmt(numbers['leave_species_out_v9_coverage'])}.",
            f"- Normalized leave-species-out known-label accuracy: {fmt(numbers['leave_species_out_v9_known_label_accuracy'])}.",
            f"- Ontology-actionable coverage: {pct(numbers['ontology_actionable_coverage'])}.",
            f"- Ontology-actionable all-cell accuracy: {pct(numbers['ontology_actionable_accuracy_all'])}.",
            f"- Ontology-label known-label accuracy: {pct(numbers['ontology_known_label_accuracy'])}.",
            f"- Ontology-label macro-F1: {fmt(numbers['ontology_macro_f1'])}.",
            f"- Adapter entries: {numbers['adapter_count']}.",
            f"- Arabidopsis root marker-candidate rows: {numbers['arabidopsis_marker_candidate_rows']}.",
            f"- Arabidopsis root cell states: {numbers['arabidopsis_cell_states']}.",
            f"- Arabidopsis root identity states: {numbers['arabidopsis_root_identity_states']}.",
            "",
            "## Claim-Safe Position",
            "",
            data["claim_safe_position"],
            "",
        ]
    )
    return "\n".join(lines)


def style_run(run: Any, size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def md_to_docx(markdown: str, output: Path) -> None:
    doc = Document()
    for line in markdown.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            paragraph = doc.add_heading(line[2:], level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                style_run(run, size=16, bold=True, color="1F4E79")
        elif line.startswith("## "):
            paragraph = doc.add_heading(line[3:], level=1)
            for run in paragraph.runs:
                style_run(run, size=13, bold=True, color="1F4E79")
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(line[2:])
            style_run(run)
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(line)
            style_run(run)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    ctx = build_context()
    synopsis = synopsis_markdown(ctx)
    highlights = highlights_json(ctx)
    highlights_md = highlights_markdown(highlights)

    MANUSCRIPT.mkdir(parents=True, exist_ok=True)
    RELEASE.mkdir(parents=True, exist_ok=True)

    synopsis_md = MANUSCRIPT / "Plant_CellFM_v9_english_submission_synopsis.md"
    synopsis_docx = MANUSCRIPT / "Plant_CellFM_v9_english_submission_synopsis.docx"
    highlights_json_path = RELEASE / "submission_highlights_v9.json"
    highlights_md_path = RELEASE / "submission_highlights_v9.md"

    synopsis_md.write_text(synopsis, encoding="utf-8")
    md_to_docx(synopsis, synopsis_docx)
    highlights_json_path.write_text(json.dumps(highlights, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    highlights_md_path.write_text(highlights_md, encoding="utf-8")

    print(synopsis_md)
    print(synopsis_docx)
    print(highlights_md_path)
    print(highlights_json_path)


if __name__ == "__main__":
    main()
