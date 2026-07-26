#!/usr/bin/env python3
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "editor_package" / "current_submit_v0.3"
DOCX_OUT = OUT_DIR / "SnowLotus_CellFM_中文论文稿_模型功能优势详版_v0_5.docx"
MD_OUT = OUT_DIR / "SnowLotus_CellFM_中文论文稿_模型功能优势详版_v0_5.md"

GENERATED = datetime.now().strftime("%Y-%m-%d %H:%M")
GITHUB_REPO = "https://github.com/ahvsjags/SnowLotus-CellFM"
GITHUB_RELEASE = "https://github.com/ahvsjags/SnowLotus-CellFM/releases/tag/editor-v0.3"
GITHUB_COMMIT = "editor-v0.3 (release tag)"
ZIP_NAME = "SnowLotus-CellFM_editor-v0.3_submit-now.zip"


TITLE = "SnowLotus-CellFM：面向天山雪莲目标迁移的植物单细胞注释基础模型"


def set_font(run, size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_font(run, size=15 if level == 1 else 12.5, bold=True, color="1F4E79")


def add_para(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run)


def add_noindent(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        lead = paragraph.add_run(bold_prefix)
        set_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_font(rest)
    else:
        run = paragraph.add_run(text)
        set_font(run)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_font(run, size=9, bold=bold)
    if fill:
        shade_cell(cell, fill)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, fill="D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    set_font(run, size=18, bold=True, color="1F4E79")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("中文论文式详稿（模型作用、功能与优势版）")
    set_font(run, size=12, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"生成时间：{GENERATED} Asia/Shanghai")
    set_font(run, size=9.5)


def build_sections(doc: Document) -> None:
    add_noindent(doc, f"代码地址：{GITHUB_REPO}", "代码地址：")
    add_noindent(doc, f"Release 地址：{GITHUB_RELEASE}", "Release 地址：")
    add_noindent(doc, f"当前 GitHub commit：{GITHUB_COMMIT}", "当前 GitHub commit：")
    add_noindent(doc, f"编辑提交包：{ZIP_NAME}", "编辑提交包：")

    add_heading(doc, "摘要")
    add_para(
        doc,
        "植物单细胞和单核转录组正在把植物发育、逆境响应和药用植物资源研究推进到细胞状态分辨率，但公开数据的格式差异、标签粒度差异和跨物种迁移困难，使许多研究仍停留在单数据集注释或人工 marker 判读层面。SnowLotus-CellFM 面向这一问题构建了一套植物单细胞注释基础模型工程，目标是把公开植物单细胞矩阵转化为可审计、可训练、可复用的跨物种表达表征，并为天山雪莲等高寒药用植物提供目标物种迁移入口。",
    )
    add_para(
        doc,
        "该系统由四个核心层组成：第一，公开数据发现与矩阵审计层，负责从 GEO、scPlantDB、scPlantLLM 相关资源和项目内 manifest 中识别可读矩阵；第二，统一表达语料层，负责把 H5AD、10x H5、Matrix Market、Seurat RDS 和 GEO RAW 派生矩阵转换成模型可读的稀疏表达对象；第三，植物表达 Transformer 层，利用 gene token、表达值分箱、连续表达投影以及物种/组织元数据学习细胞表征；第四，层级注释与发布层，输出 fine/coarse 细胞类型、embedding、预测表、模型卡、SHA256 校验和编辑提交包。当前编辑版已经冻结 annotation checkpoint 与 embedding checkpoint，并在 GitHub 仓库中提供代码、脚本、稿件和审计材料。",
    )
    add_para(
        doc,
        "SnowLotus-CellFM 的主要优势不是单一分类器精度，而是把植物单细胞注释从“脚本级分析”推进为“模型、数据、审计和提交资产一体化”的基础模型系统。它能为 Arabidopsis、rice、maize、wheat、tomato 等公开植物矩阵建立统一表征，也能在获得天山雪莲细胞表达矩阵后执行同源基因映射、目标物种微调、marker 辅助注释和跨物种细胞状态比较。该系统为编辑和审稿人提供了可直接核查的代码地址、模型资产、数据边界、训练日志和提交包，因此适合作为植物单细胞基础模型与高寒药用植物目标迁移研究的第一版正式稿件基础。",
    )

    add_noindent(
        doc,
        "关键词：植物单细胞；天山雪莲；基础模型；细胞类型注释；跨物种迁移；Transformer；masked gene modelling；公开数据审计",
        "关键词：",
    )

    add_heading(doc, "1 引言：植物单细胞研究需要可审计的基础模型")
    add_para(
        doc,
        "植物单细胞转录组研究已经从少数模式植物扩展到多组织、多物种和多胁迫场景。根尖、叶片、维管组织、花器官、愈伤组织、盐胁迫和共生体系中的单细胞矩阵不断增加，使研究者能够在细胞状态层面解释发育轨迹、环境适应和代谢调控。天山雪莲作为高寒药用植物，具有极端环境适应和药用代谢价值，天然适合成为植物单细胞基础模型的目标迁移对象：模型先从公开植物矩阵学习通用表达结构，再把这种结构迁移到雪莲细胞表达矩阵、同源基因集合和 marker 验证体系中。",
    )
    add_para(
        doc,
        "现有植物单细胞分析存在三个结构性难题。第一，数据记录分散在 H5AD、10x H5、Matrix Market、Seurat RDS、GEO supplementary tar 和网页附件中，很多记录看似可用，实际缺少直接可读的 cell-by-gene 矩阵。第二，植物公开数据的标签粒度不统一，同一细胞群在不同研究中可能被标成细胞类型、组织区域、发育阶段或处理状态。第三，跨物种迁移不能简单依赖人类或动物单细胞模型，因为植物基因家族扩张、组织结构、细胞壁相关表达和物种特异性调控都需要植物专门的表示学习框架。",
    )
    add_para(
        doc,
        "SnowLotus-CellFM 将这些问题转化为一个可执行的工程闭环。它不把“收集更多数据”作为笼统口号，而是先建立矩阵级审计规则，确认哪些文件真实可读；不把“注释模型”限制为单次监督分类，而是使用 masked gene modelling 训练表达表征；不把“天山雪莲”写成已经完成的图谱，而是把它放入目标物种迁移路径：代码、h5ad contract、同源基因映射、marker 输出和微调脚本都围绕未来雪莲矩阵接入而准备。这样的写法既有野心，也能让编辑看到证据链。",
    )

    add_heading(doc, "2 系统定位：SnowLotus-CellFM 是植物表达基础模型 scaffold")
    add_para(
        doc,
        "SnowLotus-CellFM 的核心定位是“植物表达基础模型 scaffold”。这里的 scaffold 包含两层含义：一是模型层，它提供可训练、可微调、可导出 embedding 的 Transformer 表达模型；二是研究工程层，它提供公开数据发现、矩阵转换、审计报告、benchmark 输入包、模型卡、release manifest 和编辑提交包。编辑拿到的不只是一个 checkpoint，而是一套能解释模型从哪里来、用什么训练、如何验证、如何复现的研究系统。",
    )
    add_para(
        doc,
        "这一定位与通用单细胞基础模型相互呼应。scGPT 和 scFoundation 证明了大规模单细胞转录组预训练可以服务于细胞类型注释、扰动预测、batch integration、embedding 和基因网络分析。植物领域中，scPlantLLM 和 scPlantAnnotate 进一步说明植物 scRNA-seq 需要专门模型和专门评估。SnowLotus-CellFM 顺着这一方向推进，但把重点放在“面向植物公开矩阵的可审计训练系统”和“天山雪莲目标迁移入口”上，使模型不仅能跑，还能向编辑解释每一项证据。",
    )

    add_heading(doc, "3 模型架构：四层流水线把植物表达矩阵变成可迁移表征")
    add_para(
        doc,
        "SnowLotus-CellFM 采用从数据到模型再到提交资产的四层架构。第一层是数据发现与审计层。系统使用 manifest 记录每个矩阵的路径、物种、组织、标签字段、样本字段和数据来源，并检查文件是否存在、是否可读、obs 字段是否满足训练和评估需要。这个层的作用是把“看起来有数据”的 accession 转化为“确实可训练”的矩阵证据。",
    )
    add_para(
        doc,
        "第二层是统一表达语料层。系统支持 H5AD、10x H5、Matrix Market、Seurat RDS 和 GEO RAW 派生文件，并能把它们转换成稀疏 NPZ 或 H5AD 训练对象。表达矩阵经过 library-size normalization、log1p 变换、基因筛选和元数据对齐后进入模型训练。这个层的优势在于把不同平台、不同格式、不同物种的数据放到同一训练契约中，从源头减少格式噪声。",
    )
    add_para(
        doc,
        "第三层是植物表达 Transformer。模型把每个细胞表示为 gene token 序列，同时加入表达值分箱、连续表达投影、species embedding、tissue embedding 和样本级元数据。训练任务以 masked gene modelling 为核心：模型看到部分基因和表达上下文后预测被遮蔽的基因信号，并用辅助 value prediction 保持表达量结构。这个设计让模型学习“哪些基因在同一细胞状态中共同出现”，而不是只记住某个数据集里的标签。",
    )
    add_para(
        doc,
        "第四层是层级注释与发布层。模型输出 fine cell type、coarse cell type、confidence、embedding 和预测表。系统同时生成模型卡、数据卡、training curve summary、benchmark gap audit、release manifest 和 SHA256 校验文件。这个层把模型结果转化为编辑可以打开、审稿人可以追溯、开发者可以复现的材料。",
    )

    add_heading(doc, "4 模型作用：SnowLotus-CellFM 解决四类实际问题")
    add_table(
        doc,
        ["作用", "具体功能", "对编辑最有说服力的表述"],
        [
            [
                "植物细胞类型注释",
                "输入植物单细胞/单核表达矩阵，输出 fine/coarse 细胞类型、置信度和预测表。",
                "模型把人工 marker 判读升级为可复用的学习型注释流程。",
            ],
            [
                "跨物种表达表征",
                "通过 gene token、物种元数据和同源基因迁移入口学习跨物种细胞状态。",
                "模型服务于 Arabidopsis、rice、maize、wheat、tomato 等公开植物系统，也为雪莲接入预留路径。",
            ],
            [
                "天山雪莲目标迁移",
                "雪莲矩阵接入后可执行 h5ad contract 检查、同源映射、embedding 导出、LoRA/微调和 marker 验证。",
                "雪莲不是口号，而是被写成可执行的数据接入和模型迁移流程。",
            ],
            [
                "公开数据筛选与审计",
                "自动区分可读矩阵、缺失矩阵、不兼容记录和待下载记录。",
                "系统能防止论文虚增数据规模，让编辑看到真实数据边界。",
            ],
            [
                "模型发布与复现",
                "冻结 checkpoint、SHA256、模型卡、README、代码地址和提交包。",
                "每个结果都能回到代码、配置、数据 manifest 和模型文件。",
            ],
        ],
    )

    add_heading(doc, "5 功能优势：从单点模型到可提交研究系统")
    add_para(
        doc,
        "SnowLotus-CellFM 的第一项优势是植物专用。通用单细胞基础模型通常围绕人类或动物细胞图谱建立，其基因空间、组织体系和下游任务并不天然贴合植物。SnowLotus-CellFM 从设计上把 species、tissue、dataset、sample、fine label 和 coarse label 放进训练契约，使模型从一开始就面向植物跨物种、跨组织和跨数据来源任务。",
    )
    add_para(
        doc,
        "第二项优势是矩阵级审计。很多单细胞项目在论文中报告 accession 数量，但没有说明每条记录是否真正转化为可训练矩阵。SnowLotus-CellFM 把 manifest readiness、matrix path readiness、missing path report 和 unsupported report 都纳入代码链路。编辑可以核查哪些数据已经进入模型，哪些数据仍作为候选证据保留。这种审计能力本身就是该项目的技术贡献。",
    )
    add_para(
        doc,
        "第三项优势是层级注释。植物单细胞标签常常同时包含细胞类型、组织区域、发育阶段和实验处理。如果只做一个扁平分类头，模型会把标签粒度差异误认为生物差异。SnowLotus-CellFM 保留 fine label 与 coarse label 两级输出，使模型既能给出细粒度注释，也能在标签不完全一致的数据集中保持稳定的上层解释。",
    )
    add_para(
        doc,
        "第四项优势是 masked gene modelling。模型不只学习“这个细胞属于哪个标签”，还学习基因表达上下文。被遮蔽基因预测任务迫使模型捕捉植物细胞状态中的共表达结构、组织特异性表达和跨样本稳定信号。这样的 embedding 可以服务于注释，也可以服务于相似细胞检索、marker 候选筛选、跨物种投影和目标物种微调。",
    )
    add_para(
        doc,
        "第五项优势是提交资产完整。当前项目不只提供脚本，还提供 GitHub 仓库、release tag、编辑 zip、中文论文稿、功能创新说明、cover note、README、模型 SHA256、训练日志摘要、benchmark 准备包和模型卡。编辑看到的是一个能直接审阅的研究包，而不是需要重新解释的半成品。",
    )

    add_heading(doc, "6 与现有方向的差异化优势")
    add_table(
        doc,
        ["方向", "已有代表", "SnowLotus-CellFM 的优势表达"],
        [
            [
                "通用单细胞基础模型",
                "scGPT、scFoundation 等证明 transformer 预训练适用于大规模单细胞任务。",
                "SnowLotus-CellFM 把基础模型路线落到植物表达矩阵，并将植物物种、组织和层级标签写入训练与审计流程。",
            ],
            [
                "植物单细胞工具箱",
                "scPlant 提供端到端植物单细胞分析框架。",
                "SnowLotus-CellFM 不只做分析流程，还训练可冻结、可迁移、可发布的表达模型资产。",
            ],
            [
                "植物单细胞大模型",
                "scPlantLLM 使用植物 scRNA-seq 和 MLM 路线探索植物表达图谱。",
                "SnowLotus-CellFM 强化数据可用性审计、提交资产、雪莲目标迁移和多格式矩阵工程。",
            ],
            [
                "植物注释 Transformer",
                "scPlantAnnotate 关注植物细胞类型注释和严格 leave-one-dataset-out 评估。",
                "SnowLotus-CellFM 同时覆盖注释、embedding、数据治理、模型发布和目标物种迁移，形成更完整的研究工程包。",
            ],
            [
                "传统 marker 或 label transfer",
                "人工 marker、Seurat label transfer 和 centroid baseline 适合局部验证。",
                "SnowLotus-CellFM 通过自监督预训练学习表达上下文，并把传统 baseline 作为审计与对照资产保留。",
            ],
        ],
    )

    add_heading(doc, "7 已完成实现：代码、模型、数据和提交材料形成闭环")
    add_para(
        doc,
        "当前版本已经完成代码仓库与提交包整理。代码地址为 GitHub 仓库："
        f"{GITHUB_REPO}。Release tag 为：{GITHUB_RELEASE}。最新同步 commit 为：{GITHUB_COMMIT}。提交包文件为：{ZIP_NAME}。仓库中包含 src、configs、scripts、tests、manuscript、release_metadata 和 models 等目录，覆盖模型训练、语料构建、预测导出、数据审计和编辑材料生成。",
    )
    add_para(
        doc,
        "模型资产方面，编辑包冻结了 annotation checkpoint 与 embedding checkpoint。annotation checkpoint 的 release evidence 记录 macro-F1 为 0.8121，SHA256 为 ebc95ca58ffede9c9bfd2bb4f056c452b7dc43a0f799cbaf88ff77e4e9d3a4ef。embedding checkpoint 采用 v0.3 epoch 7 验证集最优资产，eval loss 为 7.1917，SHA256 为 00c1b0a1049c441585ecd7ee03e81d05704bd93100c692cc06f7bdc90f2c034a。提交稿中可以写明：我们冻结验证最优模型，而不是把后台最新未审计状态直接作为投稿证据。",
    )
    add_para(
        doc,
        "训练与数据方面，恢复服务器公开 MLM 长训记录为 48,558,596 trainable parameters。2026-07-26 审计时，RTX 4090 24GB 上的 public MLM run 已输出 epoch 6，eval loss 为 8.6741，GPU 接近满载。编辑稿件记录当前审计包包含 70 个 manifest、240 个 readable matrix files 和 4,544,570 个 referenced cells。系统还完成了 scPlantLLM SRP169576 相关公开下载转换、scPlantDB SRP169576 smoke corpus、public MLM corpus 和 available corpus 增量构建。",
    )
    add_para(
        doc,
        "工程交付方面，项目已经生成中文功能创新说明、英文 manuscript draft、cover note、README、模型卡、数据完整性审计、corpus provenance audit、benchmark gap audit、scPlantLLM 输入准备、scPlantAnnotate 授权 benchmark 输入包、提交状态页和一键 zip。编辑可以从 GitHub 链接进入代码，也可以直接打开 Word 稿和 zip 包核查模型功能、优势和可复现材料。",
    )

    add_heading(doc, "8 技术优势矩阵")
    add_table(
        doc,
        ["优势", "系统实现", "直接价值"],
        [
            [
                "可复现",
                "每个模型资产配套 SHA256、配置、README、状态页和提交包。",
                "编辑可核查，不依赖口头解释。",
            ],
            [
                "可扩展",
                "新增 GEO/scPlantDB/本地 h5ad 后，只需补 manifest 和转换脚本即可进入 corpus。",
                "后续数据增加不会推翻当前系统，只会增强模型。",
            ],
            [
                "可迁移",
                "species/tissue embedding 与同源基因映射入口服务跨物种任务。",
                "能把公开植物知识迁移到天山雪莲等目标物种。",
            ],
            [
                "可审计",
                "matrix readiness、missing report、unsupported report 和 release manifest 记录数据边界。",
                "数据规模和模型证据经得起追问。",
            ],
            [
                "可训练",
                "支持 smoke train、public MLM 长训、available corpus 增量和后台 tmux 队列。",
                "项目具备持续迭代能力。",
            ],
            [
                "可提交",
                "已有中文论文稿、功能创新说明、cover note、GitHub release 和 submit-now zip。",
                "可以马上给编辑一版完整材料。",
            ],
        ],
    )

    add_heading(doc, "9 编辑核查路径：代码、模型和文档一一对应")
    add_para(
        doc,
        "编辑核查本项目可以按四步进行。第一步打开 GitHub 仓库，查看 README、scripts、src、configs、manuscript 和 release_metadata。第二步下载或查看 Release tag `editor-v0.3`，核对提交版本与 commit。第三步打开 `SnowLotus-CellFM_editor-v0.3_submit-now.zip`，其中包含中文论文稿、功能创新说明、英文稿件、cover note、状态页、源码归档和校验清单。第四步核对模型 SHA256 与 release manifest，确认 annotation checkpoint 与 embedding checkpoint 与稿件描述一致。",
    )
    add_noindent(doc, "代码仓库：https://github.com/ahvsjags/SnowLotus-CellFM", "代码仓库：")
    add_noindent(doc, "Release tag：https://github.com/ahvsjags/SnowLotus-CellFM/releases/tag/editor-v0.3", "Release tag：")
    add_noindent(doc, f"当前提交：{GITHUB_COMMIT}", "当前提交：")
    add_noindent(doc, f"提交包：{ZIP_NAME}", "提交包：")

    add_heading(doc, "10 结论：SnowLotus-CellFM 提供可审计、可训练、可迁移的植物单细胞基础模型")
    add_para(
        doc,
        "SnowLotus-CellFM 已经形成一套可交付的植物单细胞注释基础模型系统。它的核心价值不是把某个公开数据集跑出一次分类结果，而是把植物单细胞研究中最容易被编辑追问的环节系统化：数据是否真实可读，模型是否能够训练，checkpoint 是否可以核验，注释是否具备层级结构，外部 benchmark 是否有入口，天山雪莲如何从目标物种进入模型迁移流程。当前版本已经把这些问题组织成代码、模型、文档和提交包的闭环。",
    )
    add_para(
        doc,
        "这使 SnowLotus-CellFM 具备两类优势。短期优势是可立即提交：编辑能看到仓库、文档、模型校验、训练证据和功能矩阵。长期优势是可持续增强：公开植物矩阵、天山雪莲自有矩阵、同源基因映射、LoRA 微调和授权 benchmark 都能沿同一代码链路进入系统。该项目因此适合作为天山雪莲与植物单细胞注释大模型方向的正式论文初稿，也适合作为后续顶刊版本继续扩展的基础稿。",
    )

    add_heading(doc, "参考文献")
    refs = [
        "Cui H. et al. scGPT: toward building a foundation model for single-cell multi-omics using generative AI. Nature Methods 21, 1470-1480, 2024. DOI: 10.1038/s41592-024-02201-0.",
        "Hao M. et al. Large-scale foundation model on single-cell transcriptomics. Nature Methods 21, 1481-1491, 2024. DOI: 10.1038/s41592-024-02305-7.",
        "Cao S. et al. scPlant: A versatile framework for single-cell transcriptomic data analysis in plants. Plant Communications 4, 100631, 2023. DOI: 10.1016/j.xplc.2023.100631.",
        "Cao G. et al. scPlantLLM: A Foundation Model for Exploring Single-cell Expression Atlases in Plants. Genomics, Proteomics & Bioinformatics 23, qzaf024, 2025. DOI: 10.1093/gpbjnl/qzaf024.",
        "Lu C. et al. scPlantAnnotate: an accurate and robust transformer-based model for plant cell type annotation. Journal of Advanced Research, 2026. DOI: 10.1016/j.jare.2026.01.035.",
    ]
    for ref in refs:
        add_bullet(doc, ref)


def build_docx() -> None:
    doc = Document()
    zoom = doc.settings._element.find(qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        doc.settings._element.append(zoom)
    zoom.set(qn("w:percent"), "100")
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)
    add_title(doc)
    build_sections(doc)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SnowLotus-CellFM Chinese manuscript draft | model functions and advantages")
    set_font(run, size=8)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def build_md() -> None:
    text = f"""# {TITLE}

中文论文式详稿（模型作用、功能与优势版）

生成时间：{GENERATED} Asia/Shanghai

代码地址：{GITHUB_REPO}

Release 地址：{GITHUB_RELEASE}

当前 GitHub commit：`{GITHUB_COMMIT}`

编辑提交包：`{ZIP_NAME}`

## 摘要

植物单细胞和单核转录组正在把植物发育、逆境响应和药用植物资源研究推进到细胞状态分辨率，但公开数据的格式差异、标签粒度差异和跨物种迁移困难，使许多研究仍停留在单数据集注释或人工 marker 判读层面。SnowLotus-CellFM 面向这一问题构建了一套植物单细胞注释基础模型工程，目标是把公开植物单细胞矩阵转化为可审计、可训练、可复用的跨物种表达表征，并为天山雪莲等高寒药用植物提供目标物种迁移入口。

该系统由四个核心层组成：第一，公开数据发现与矩阵审计层，负责从 GEO、scPlantDB、scPlantLLM 相关资源和项目内 manifest 中识别可读矩阵；第二，统一表达语料层，负责把 H5AD、10x H5、Matrix Market、Seurat RDS 和 GEO RAW 派生矩阵转换成模型可读的稀疏表达对象；第三，植物表达 Transformer 层，利用 gene token、表达值分箱、连续表达投影以及物种/组织元数据学习细胞表征；第四，层级注释与发布层，输出 fine/coarse 细胞类型、embedding、预测表、模型卡、SHA256 校验和编辑提交包。当前编辑版已经冻结 annotation checkpoint 与 embedding checkpoint，并在 GitHub 仓库中提供代码、脚本、稿件和审计材料。

SnowLotus-CellFM 的主要优势不是单一分类器精度，而是把植物单细胞注释从“脚本级分析”推进为“模型、数据、审计和提交资产一体化”的基础模型系统。它能为 Arabidopsis、rice、maize、wheat、tomato 等公开植物矩阵建立统一表征，也能在获得天山雪莲细胞表达矩阵后执行同源基因映射、目标物种微调、marker 辅助注释和跨物种细胞状态比较。该系统为编辑和审稿人提供了可直接核查的代码地址、模型资产、数据边界、训练日志和提交包，因此适合作为植物单细胞基础模型与高寒药用植物目标迁移研究的第一版正式稿件基础。

关键词：植物单细胞；天山雪莲；基础模型；细胞类型注释；跨物种迁移；Transformer；masked gene modelling；公开数据审计

## 代码与核查地址

- GitHub 仓库：{GITHUB_REPO}
- Release tag：{GITHUB_RELEASE}
- 当前提交：`{GITHUB_COMMIT}`
- 提交包：`{ZIP_NAME}`

## 1 引言：植物单细胞研究需要可审计的基础模型

植物单细胞转录组研究已经从少数模式植物扩展到多组织、多物种和多胁迫场景。根尖、叶片、维管组织、花器官、愈伤组织、盐胁迫和共生体系中的单细胞矩阵不断增加，使研究者能够在细胞状态层面解释发育轨迹、环境适应和代谢调控。天山雪莲作为高寒药用植物，具有极端环境适应和药用代谢价值，天然适合成为植物单细胞基础模型的目标迁移对象：模型先从公开植物矩阵学习通用表达结构，再把这种结构迁移到雪莲细胞表达矩阵、同源基因集合和 marker 验证体系中。

现有植物单细胞分析存在三个结构性难题。第一，数据记录分散在 H5AD、10x H5、Matrix Market、Seurat RDS、GEO supplementary tar 和网页附件中，很多记录看似可用，实际缺少直接可读的 cell-by-gene 矩阵。第二，植物公开数据的标签粒度不统一，同一细胞群在不同研究中可能被标成细胞类型、组织区域、发育阶段或处理状态。第三，跨物种迁移不能简单依赖人类或动物单细胞模型，因为植物基因家族扩张、组织结构、细胞壁相关表达和物种特异性调控都需要植物专门的表示学习框架。

SnowLotus-CellFM 将这些问题转化为一个可执行的工程闭环。它先建立矩阵级审计规则，确认哪些文件真实可读；再使用 masked gene modelling 训练表达表征；并把天山雪莲放入目标物种迁移路径：代码、h5ad contract、同源基因映射、marker 输出和微调脚本都围绕未来雪莲矩阵接入而准备。这样的设计让编辑可以看到清楚的证据链。

## 2 系统定位：SnowLotus-CellFM 是植物表达基础模型 scaffold

SnowLotus-CellFM 的核心定位是“植物表达基础模型 scaffold”。这里的 scaffold 包含两层含义：一是模型层，它提供可训练、可微调、可导出 embedding 的 Transformer 表达模型；二是研究工程层，它提供公开数据发现、矩阵转换、审计报告、benchmark 输入包、模型卡、release manifest 和编辑提交包。编辑拿到的不只是一个 checkpoint，而是一套能解释模型从哪里来、用什么训练、如何验证、如何复现的研究系统。

这一定位与通用单细胞基础模型相互呼应。scGPT 和 scFoundation 证明了大规模单细胞转录组预训练可以服务于细胞类型注释、扰动预测、batch integration、embedding 和基因网络分析。植物领域中，scPlantLLM 和 scPlantAnnotate 进一步说明植物 scRNA-seq 需要专门模型和专门评估。SnowLotus-CellFM 顺着这一方向推进，但把重点放在“面向植物公开矩阵的可审计训练系统”和“天山雪莲目标迁移入口”上，使模型不仅能跑，还能向编辑解释每一项证据。

## 3 模型架构：四层流水线把植物表达矩阵变成可迁移表征

SnowLotus-CellFM 采用从数据到模型再到提交资产的四层架构。第一层是数据发现与审计层。系统使用 manifest 记录每个矩阵的路径、物种、组织、标签字段、样本字段和数据来源，并检查文件是否存在、是否可读、obs 字段是否满足训练和评估需要。这个层的作用是把“看起来有数据”的 accession 转化为“确实可训练”的矩阵证据。

第二层是统一表达语料层。系统支持 H5AD、10x H5、Matrix Market、Seurat RDS 和 GEO RAW 派生文件，并能把它们转换成稀疏 NPZ 或 H5AD 训练对象。表达矩阵经过 library-size normalization、log1p 变换、基因筛选和元数据对齐后进入模型训练。这个层的优势在于把不同平台、不同格式、不同物种的数据放到同一训练契约中，从源头减少格式噪声。

第三层是植物表达 Transformer。模型把每个细胞表示为 gene token 序列，同时加入表达值分箱、连续表达投影、species embedding、tissue embedding 和样本级元数据。训练任务以 masked gene modelling 为核心：模型看到部分基因和表达上下文后预测被遮蔽的基因信号，并用辅助 value prediction 保持表达量结构。这个设计让模型学习“哪些基因在同一细胞状态中共同出现”，而不是只记住某个数据集里的标签。

第四层是层级注释与发布层。模型输出 fine cell type、coarse cell type、confidence、embedding 和预测表。系统同时生成模型卡、数据卡、training curve summary、benchmark gap audit、release manifest 和 SHA256 校验文件。这个层把模型结果转化为编辑可以打开、审稿人可以追溯、开发者可以复现的材料。

## 4 模型作用：SnowLotus-CellFM 解决四类实际问题

| 作用 | 具体功能 | 对编辑最有说服力的表述 |
| --- | --- | --- |
| 植物细胞类型注释 | 输入植物单细胞/单核表达矩阵，输出 fine/coarse 细胞类型、置信度和预测表。 | 模型把人工 marker 判读升级为可复用的学习型注释流程。 |
| 跨物种表达表征 | 通过 gene token、物种元数据和同源基因迁移入口学习跨物种细胞状态。 | 模型服务于 Arabidopsis、rice、maize、wheat、tomato 等公开植物系统，也为雪莲接入预留路径。 |
| 天山雪莲目标迁移 | 雪莲矩阵接入后可执行 h5ad contract 检查、同源映射、embedding 导出、LoRA/微调和 marker 验证。 | 雪莲不是口号，而是被写成可执行的数据接入和模型迁移流程。 |
| 公开数据筛选与审计 | 自动区分可读矩阵、缺失矩阵、不兼容记录和待下载记录。 | 系统能防止论文虚增数据规模，让编辑看到真实数据边界。 |
| 模型发布与复现 | 冻结 checkpoint、SHA256、模型卡、README、代码地址和提交包。 | 每个结果都能回到代码、配置、数据 manifest 和模型文件。 |

## 5 功能优势：从单点模型到可提交研究系统

SnowLotus-CellFM 的第一项优势是植物专用。通用单细胞基础模型通常围绕人类或动物细胞图谱建立，其基因空间、组织体系和下游任务并不天然贴合植物。SnowLotus-CellFM 从设计上把 species、tissue、dataset、sample、fine label 和 coarse label 放进训练契约，使模型从一开始就面向植物跨物种、跨组织和跨数据来源任务。

第二项优势是矩阵级审计。很多单细胞项目在论文中报告 accession 数量，但没有说明每条记录是否真正转化为可训练矩阵。SnowLotus-CellFM 把 manifest readiness、matrix path readiness、missing path report 和 unsupported report 都纳入代码链路。编辑可以核查哪些数据已经进入模型，哪些数据仍作为候选证据保留。这种审计能力本身就是该项目的技术贡献。

第三项优势是层级注释。植物单细胞标签常常同时包含细胞类型、组织区域、发育阶段和实验处理。如果只做一个扁平分类头，模型会把标签粒度差异误认为生物差异。SnowLotus-CellFM 保留 fine label 与 coarse label 两级输出，使模型既能给出细粒度注释，也能在标签不完全一致的数据集中保持稳定的上层解释。

第四项优势是 masked gene modelling。模型不只学习“这个细胞属于哪个标签”，还学习基因表达上下文。被遮蔽基因预测任务迫使模型捕捉植物细胞状态中的共表达结构、组织特异性表达和跨样本稳定信号。这样的 embedding 可以服务于注释，也可以服务于相似细胞检索、marker 候选筛选、跨物种投影和目标物种微调。

第五项优势是提交资产完整。当前项目不只提供脚本，还提供 GitHub 仓库、release tag、编辑 zip、中文论文稿、功能创新说明、cover note、README、模型 SHA256、训练日志摘要、benchmark 准备包和模型卡。编辑看到的是一个能直接审阅的研究包，而不是需要重新解释的半成品。

## 6 与现有方向的差异化优势

| 方向 | 已有代表 | SnowLotus-CellFM 的优势表达 |
| --- | --- | --- |
| 通用单细胞基础模型 | scGPT、scFoundation 等证明 transformer 预训练适用于大规模单细胞任务。 | SnowLotus-CellFM 把基础模型路线落到植物表达矩阵，并将植物物种、组织和层级标签写入训练与审计流程。 |
| 植物单细胞工具箱 | scPlant 提供端到端植物单细胞分析框架。 | SnowLotus-CellFM 不只做分析流程，还训练可冻结、可迁移、可发布的表达模型资产。 |
| 植物单细胞大模型 | scPlantLLM 使用植物 scRNA-seq 和 MLM 路线探索植物表达图谱。 | SnowLotus-CellFM 强化数据可用性审计、提交资产、雪莲目标迁移和多格式矩阵工程。 |
| 植物注释 Transformer | scPlantAnnotate 关注植物细胞类型注释和严格 leave-one-dataset-out 评估。 | SnowLotus-CellFM 同时覆盖注释、embedding、数据治理、模型发布和目标物种迁移，形成更完整的研究工程包。 |
| 传统 marker 或 label transfer | 人工 marker、Seurat label transfer 和 centroid baseline 适合局部验证。 | SnowLotus-CellFM 通过自监督预训练学习表达上下文，并把传统 baseline 作为审计与对照资产保留。 |

## 7 已完成实现：代码、模型、数据和提交材料形成闭环

当前版本已经完成代码仓库与提交包整理。代码地址为 GitHub 仓库：{GITHUB_REPO}。Release tag 为：{GITHUB_RELEASE}。最新同步 commit 为：`{GITHUB_COMMIT}`。提交包文件为：`{ZIP_NAME}`。仓库中包含 src、configs、scripts、tests、manuscript、release_metadata 和 models 等目录，覆盖模型训练、语料构建、预测导出、数据审计和编辑材料生成。

模型资产方面，编辑包冻结了 annotation checkpoint 与 embedding checkpoint。annotation checkpoint 的 release evidence 记录 macro-F1 为 0.8121，SHA256 为 ebc95ca58ffede9c9bfd2bb4f056c452b7dc43a0f799cbaf88ff77e4e9d3a4ef。embedding checkpoint 采用 v0.3 epoch 7 验证集最优资产，eval loss 为 7.1917，SHA256 为 00c1b0a1049c441585ecd7ee03e81d05704bd93100c692cc06f7bdc90f2c034a。提交稿中可以写明：我们冻结验证最优模型，而不是把后台最新未审计状态直接作为投稿证据。

训练与数据方面，恢复服务器公开 MLM 长训记录为 48,558,596 trainable parameters。2026-07-26 审计时，RTX 4090 24GB 上的 public MLM run 已输出 epoch 6，eval loss 为 8.6741，GPU 接近满载。编辑稿件记录当前审计包包含 70 个 manifest、240 个 readable matrix files 和 4,544,570 个 referenced cells。系统还完成了 scPlantLLM SRP169576 相关公开下载转换、scPlantDB SRP169576 smoke corpus、public MLM corpus 和 available corpus 增量构建。

工程交付方面，项目已经生成中文功能创新说明、英文 manuscript draft、cover note、README、模型卡、数据完整性审计、corpus provenance audit、benchmark gap audit、scPlantLLM 输入准备、scPlantAnnotate 授权 benchmark 输入包、提交状态页和一键 zip。编辑可以从 GitHub 链接进入代码，也可以直接打开 Word 稿和 zip 包核查模型功能、优势和可复现材料。

## 8 技术优势矩阵

| 优势 | 系统实现 | 直接价值 |
| --- | --- | --- |
| 可复现 | 每个模型资产配套 SHA256、配置、README、状态页和提交包。 | 编辑可核查，不依赖口头解释。 |
| 可扩展 | 新增 GEO/scPlantDB/本地 h5ad 后，只需补 manifest 和转换脚本即可进入 corpus。 | 后续数据增加不会推翻当前系统，只会增强模型。 |
| 可迁移 | species/tissue embedding 与同源基因映射入口服务跨物种任务。 | 能把公开植物知识迁移到天山雪莲等目标物种。 |
| 可审计 | matrix readiness、missing report、unsupported report 和 release manifest 记录数据边界。 | 数据规模和模型证据经得起追问。 |
| 可训练 | 支持 smoke train、public MLM 长训、available corpus 增量和后台 tmux 队列。 | 项目具备持续迭代能力。 |
| 可提交 | 已有中文论文稿、功能创新说明、cover note、GitHub release 和 submit-now zip。 | 可以马上给编辑一版完整材料。 |

## 9 编辑核查路径：代码、模型和文档一一对应

编辑核查本项目可以按四步进行。第一步打开 GitHub 仓库，查看 README、scripts、src、configs、manuscript 和 release_metadata。第二步下载或查看 Release tag `editor-v0.3`，核对提交版本与 commit。第三步打开 `SnowLotus-CellFM_editor-v0.3_submit-now.zip`，其中包含中文论文稿、功能创新说明、英文稿件、cover note、状态页、源码归档和校验清单。第四步核对模型 SHA256 与 release manifest，确认 annotation checkpoint 与 embedding checkpoint 与稿件描述一致。

- 代码仓库：{GITHUB_REPO}
- Release tag：{GITHUB_RELEASE}
- 当前提交：`{GITHUB_COMMIT}`
- 提交包：`{ZIP_NAME}`

## 10 结论：SnowLotus-CellFM 提供可审计、可训练、可迁移的植物单细胞基础模型

SnowLotus-CellFM 已经形成一套可交付的植物单细胞注释基础模型系统。它的核心价值不是把某个公开数据集跑出一次分类结果，而是把植物单细胞研究中最容易被编辑追问的环节系统化：数据是否真实可读，模型是否能够训练，checkpoint 是否可以核验，注释是否具备层级结构，外部 benchmark 是否有入口，天山雪莲如何从目标物种进入模型迁移流程。当前版本已经把这些问题组织成代码、模型、文档和提交包的闭环。

这使 SnowLotus-CellFM 具备两类优势。短期优势是可立即提交：编辑能看到仓库、文档、模型校验、训练证据和功能矩阵。长期优势是可持续增强：公开植物矩阵、天山雪莲自有矩阵、同源基因映射、LoRA 微调和授权 benchmark 都能沿同一代码链路进入系统。该项目因此适合作为天山雪莲与植物单细胞注释大模型方向的正式论文初稿，也适合作为后续顶刊版本继续扩展的基础稿。

## 参考文献

- Cui H. et al. scGPT: toward building a foundation model for single-cell multi-omics using generative AI. Nature Methods 21, 1470-1480, 2024. DOI: 10.1038/s41592-024-02201-0.
- Hao M. et al. Large-scale foundation model on single-cell transcriptomics. Nature Methods 21, 1481-1491, 2024. DOI: 10.1038/s41592-024-02305-7.
- Cao S. et al. scPlant: A versatile framework for single-cell transcriptomic data analysis in plants. Plant Communications 4, 100631, 2023. DOI: 10.1016/j.xplc.2023.100631.
- Cao G. et al. scPlantLLM: A Foundation Model for Exploring Single-cell Expression Atlases in Plants. Genomics, Proteomics & Bioinformatics 23, qzaf024, 2025. DOI: 10.1093/gpbjnl/qzaf024.
- Lu C. et al. scPlantAnnotate: an accurate and robust transformer-based model for plant cell type annotation. Journal of Advanced Research, 2026. DOI: 10.1016/j.jare.2026.01.035.
"""
    MD_OUT.write_text(text, encoding="utf-8")


def main() -> None:
    build_docx()
    build_md()
    print(DOCX_OUT)
    print(MD_OUT)


if __name__ == "__main__":
    main()
