#!/usr/bin/env python3
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "editor_package" / "current_submit_v0.3"
DOCX_OUT = OUT_DIR / "SnowLotus_CellFM_已完成工作校稿版_v0_4.docx"
MD_OUT = OUT_DIR / "SnowLotus_CellFM_已完成工作校稿版_v0_4.md"

GENERATED = datetime.now().strftime("%Y-%m-%d %H:%M")
GITHUB_REPO = "https://github.com/ahvsjags/SnowLotus-CellFM"
GITHUB_RELEASE = "https://github.com/ahvsjags/SnowLotus-CellFM/releases/tag/editor-v0.3"
LATEST_PUSHED_COMMIT = "099e10081ededdedf025513350736733decfed09"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_para(doc: Document, text: str, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        lead.bold = True
        rest = paragraph.add_run(text[len(bold_lead) :])
        runs = [lead, rest]
    else:
        runs = [paragraph.add_run(text)]
    for run in runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SnowLotus-CellFM\n天山雪莲与植物单细胞注释大模型\n已完成工作汇总（校稿版）")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"生成时间：{GENERATED} Asia/Shanghai")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)

    add_para(doc, f"GitHub 仓库：{GITHUB_REPO}", "GitHub 仓库：")
    add_para(doc, f"Release 标签：{GITHUB_RELEASE}", "Release 标签：")
    add_para(doc, f"当前已推送脚本修复 commit：{LATEST_PUSHED_COMMIT}", "当前已推送脚本修复 commit：")
    add_para(
        doc,
        "说明：本文件只整理已经完成、可供编辑或合作者校稿的内容；仍在下载、训练、授权 benchmark 或远程端口恢复中的事项，均放在“边界与待补强”中，不作为已完成结果表述。",
        "说明：",
    )

    add_heading(doc, "一、项目一句话定位")
    add_para(
        doc,
        "SnowLotus-CellFM 是一个面向植物跨物种单细胞/单核表达数据的注释与表征基础模型工程。当前版本已经完成从公开数据发现、矩阵审计、语料构建、Transformer 训练、层级注释、外部 benchmark 准备、模型冻结、GitHub 发布到编辑提交包整理的一条可复现链路；天山雪莲被明确定位为高寒药用植物的目标迁移与后续实验验证场景，而不是夸大为已经完成的雪莲单细胞图谱。",
    )

    add_heading(doc, "二、现在可以校稿的核心结论")
    for item in [
        "项目不是停留在方案阶段，已经形成可运行代码仓库、训练脚本、数据审计脚本、模型 checkpoint、投稿说明和一键提交包。",
        "当前可展示的模型资产包括冻结 annotation checkpoint 与 embedding checkpoint，并有 SHA256 校验、模型卡、release manifest 和编辑包记录。",
        "公开植物单细胞数据链路已经能处理 H5AD、10x H5、Matrix Market、Seurat RDS、GEO RAW tar 等多种格式，并能区分可用矩阵、缺失矩阵和不兼容记录。",
        "模型路线采用植物表达 gene-token / expression-value / species-tissue metadata 的 Transformer masked modelling，并保留 fine/coarse 层级注释头，适合跨物种迁移和目标植物适配。",
        "天山雪莲部分当前应写作“目标物种迁移框架与数据缺口已定义”，不应写作“雪莲单细胞图谱已经完成”。这一点对编辑和审稿人更可信。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "三、已完成工作总览")
    add_table(
        doc,
        ["模块", "已完成内容", "可校稿写法"],
        [
            [
                "项目与仓库",
                "已建立 SnowLotus-CellFM 代码仓库、编辑提交包、中文说明、英文 manuscript draft、cover note、README、release tag。",
                "项目具备可复现代码与编辑审阅材料，不再只是概念方案。",
            ],
            [
                "数据链路",
                "已完成公开植物单细胞数据发现、manifest 管理、矩阵完整性审计、格式转换和可用/不可用边界记录。",
                "模型训练语料经过矩阵级审计，避免把不可读或仅元数据记录虚增为训练数据。",
            ],
            [
                "模型训练",
                "已完成 smoke train/predict、监督注释 checkpoint、公开 MLM embedding checkpoint，并在恢复服务器上继续长训。",
                "当前版本可展示冻结模型资产，同时后台继续改进更强 checkpoint。",
            ],
            [
                "评价与审计",
                "已整理 strict split audit、centroid baseline、Seurat/scPlantLLM 输入路径、scPlantAnnotate 授权 benchmark 准备包、模型卡和 release manifest。",
                "评价体系包含内部指标与外部工具对标准备，缺口被显式列出。",
            ],
            [
                "编辑提交",
                "已生成 editor-v0.3 zip、中文功能创新说明、稿件草稿、cover note、README、状态页和 GitHub 链接。",
                "可先给编辑一版能看懂创新与完成度的材料，后续再用新训练结果补强。",
            ],
        ],
    )

    add_heading(doc, "四、可写进稿件的创新点")
    add_table(
        doc,
        ["创新点", "具体含义", "为什么适合顶刊叙事"],
        [
            [
                "植物单细胞基础模型工程化",
                "把植物单细胞注释从单个分类脚本升级为可持续训练、可审计、可发布的 foundation-model scaffold。",
                "解决植物公开单细胞数据碎片化、复现难和跨物种泛化难的问题。",
            ],
            [
                "矩阵级数据审计前置",
                "训练前逐条检查矩阵文件是否存在、可读、字段完整，并保留 unsupported/deferred 记录。",
                "让数据规模和训练证据更可信，避免审稿人质疑数据虚增。",
            ],
            [
                "gene-token + expression-value Transformer",
                "每个细胞由高表达基因、表达值分箱、连续表达投影和物种/组织元数据共同表征。",
                "比普通 marker 规则更适合跨平台、跨物种、跨组织表达结构学习。",
            ],
            [
                "层级细胞类型注释",
                "同时保留 fine label 和 coarse label，缓解不同数据集注释粒度不一致的问题。",
                "植物单细胞公开数据标签体系混乱，这一点很容易打动审稿人。",
            ],
            [
                "天山雪莲目标迁移路线",
                "把雪莲定位为后续适配对象，准备 h5ad contract、同源基因映射、marker 验证和 LoRA/微调路径。",
                "诚实处理雪莲公开单细胞矩阵缺口，同时形成清楚的实验转化路线。",
            ],
            [
                "训练与提交并行",
                "编辑版冻结可提交证据，服务器后台继续下载公开数据和训练新模型，后续用同一审计链路更新。",
                "适合赶编辑时先提交一版，同时保留下一轮增强空间。",
            ],
        ],
    )

    add_heading(doc, "五、当前可引用的关键证据")
    add_table(
        doc,
        ["证据类型", "当前事实", "建议写法"],
        [
            [
                "模型规模",
                "恢复服务器公开 MLM 长训记录为 48,558,596 trainable parameters。",
                "当前模型为约 4,856 万可训练参数的植物表达 Transformer。",
            ],
            [
                "Embedding checkpoint",
                "编辑包冻结 v0.3 epoch 7 embedding asset，eval loss 7.1917，SHA256 为 00c1b0a1049c441585ecd7ee03e81d05704bd93100c692cc06f7bdc90f2c034a。",
                "冻结提交模型采用验证集最优 checkpoint，而不是后台最新未审计状态。",
            ],
            [
                "Annotation checkpoint",
                "冻结 annotation checkpoint 的 release evidence 记录 macro-F1 0.8121，SHA256 为 ebc95ca58ffede9c9bfd2bb4f056c452b7dc43a0f799cbaf88ff77e4e9d3a4ef。",
                "注释模型已有可引用监督指标，后续可补更严格外部 benchmark。",
            ],
            [
                "恢复服务器训练",
                "2026-07-26 审计时，RTX 4090 24GB 上 public MLM run 已到 epoch 6，eval loss 8.6741，GPU 接近满载。",
                "远程训练链路已经跑通，服务器仍作为后台改进通道。",
            ],
            [
                "公开语料",
                "编辑稿件记录当前审计包包含 70 manifest、240 readable matrix files、4,544,570 referenced cells；恢复服务器另有 reconstructed public MLM corpus 及后续 available corpus 增量。",
                "公开数据语料按可读矩阵审计，不将未下载或不可读记录混入完成证据。",
            ],
            [
                "提交包",
                "本地已有 SnowLotus-CellFM_editor-v0.3_submit-now.zip、中文功能创新说明、稿件草稿、cover note、README 和模型校验信息。",
                "可以先作为编辑沟通版提交，后续 revision 再补充更大 corpus 和 benchmark。",
            ],
        ],
    )

    add_heading(doc, "六、建议给编辑/合作者看的摘要版本")
    add_para(
        doc,
        "我们已经完成 SnowLotus-CellFM 的第一版可复现研究发布。该项目面向植物单细胞和单核转录组数据，建立了从公开数据发现、矩阵级审计、语料构建、Transformer masked modelling、层级细胞类型注释、模型冻结、外部 benchmark 准备到 GitHub/编辑包交付的完整链路。当前版本冻结了可供审阅的 annotation 与 embedding checkpoint，保留 SHA256 校验和模型卡，并将天山雪莲明确作为目标物种迁移场景处理。现阶段不夸大为已经完成雪莲单细胞图谱，而是提供一个可立即审阅、可继续训练、可在获得雪莲单细胞矩阵后快速适配的植物单细胞基础模型框架。",
    )

    add_heading(doc, "七、边界与待补强")
    add_table(
        doc,
        ["事项", "当前边界", "下一步"],
        [
            [
                "天山雪莲单细胞矩阵",
                "当前审计未确认可复用的公开 Snow Lotus scRNA/snRNA cell-by-gene matrix。",
                "继续按数据请求包或自有实验接入 h5ad/10x/MTX 矩阵，再做目标物种微调和注释验证。",
            ],
            [
                "远程服务器状态",
                "2026-07-26 已验证训练在 Matpool 恢复服务器运行；2026-07-27 当前端口返回 Connection refused，说明实例端口需恢复或更换。",
                "服务器恢复后同步本次 pipeline 修复脚本，继续数据队列和训练。",
            ],
            [
                "外部 benchmark",
                "scPlantAnnotate 需要授权或结果导出；当前已准备输入包和访问审计，但不能写成完成指标。",
                "拿到授权后运行预测并写入 final metrics。",
            ],
            [
                "GitHub 可见性",
                "仓库当前按状态页记录为 private。",
                "给编辑/审稿人使用前，需要添加访问权限或切换为 public。",
            ],
        ],
    )

    add_heading(doc, "八、校稿时重点看哪里")
    for item in [
        "标题是否要更偏“技术资源”还是更偏“天山雪莲应用”。",
        "是否接受把天山雪莲写成 target-species transfer，而不是完成 atlas。",
        "模型指标表述是否保守：冻结 checkpoint 与后台训练状态要分开。",
        "公开数据规模是否按“审计包记录”和“恢复服务器实时增量”分开写。",
        "GitHub 链接在提交前是否需要公开或授权给编辑。",
    ]:
        add_bullet(doc, item)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SnowLotus-CellFM completed-work proof draft | editor-v0.3/v0.4 working note")
    run.font.size = Pt(8)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_OUT)


def build_md() -> None:
    text = f"""# SnowLotus-CellFM 已完成工作汇总（校稿版）

生成时间：{GENERATED} Asia/Shanghai

GitHub 仓库：{GITHUB_REPO}

Release 标签：{GITHUB_RELEASE}

当前已推送脚本修复 commit：`{LATEST_PUSHED_COMMIT}`

本文件只整理已经完成、可供编辑或合作者校稿的内容；仍在下载、训练、授权 benchmark 或远程端口恢复中的事项，均放在边界与待补强中，不作为已完成结果表述。

## 一句话定位

SnowLotus-CellFM 是一个面向植物跨物种单细胞/单核表达数据的注释与表征基础模型工程。当前版本已经完成从公开数据发现、矩阵审计、语料构建、Transformer 训练、层级注释、外部 benchmark 准备、模型冻结、GitHub 发布到编辑提交包整理的一条可复现链路；天山雪莲被明确定位为高寒药用植物的目标迁移与后续实验验证场景，而不是夸大为已经完成的雪莲单细胞图谱。

## 已完成且可校稿的核心结论

- 项目已经形成可运行代码仓库、训练脚本、数据审计脚本、模型 checkpoint、投稿说明和一键提交包。
- 当前可展示的模型资产包括冻结 annotation checkpoint 与 embedding checkpoint，并有 SHA256 校验、模型卡、release manifest 和编辑包记录。
- 公开植物单细胞数据链路已经能处理 H5AD、10x H5、Matrix Market、Seurat RDS、GEO RAW tar 等多种格式，并能区分可用矩阵、缺失矩阵和不兼容记录。
- 模型路线采用植物表达 gene-token / expression-value / species-tissue metadata 的 Transformer masked modelling，并保留 fine/coarse 层级注释头。
- 天山雪莲部分当前应写作“目标物种迁移框架与数据缺口已定义”，不应写作“雪莲单细胞图谱已经完成”。

## 可引用证据

- 模型规模：恢复服务器公开 MLM 长训记录为 48,558,596 trainable parameters。
- Embedding checkpoint：v0.3 epoch 7，eval loss 7.1917，SHA256 `00c1b0a1049c441585ecd7ee03e81d05704bd93100c692cc06f7bdc90f2c034a`。
- Annotation checkpoint：release evidence 记录 macro-F1 0.8121，SHA256 `ebc95ca58ffede9c9bfd2bb4f056c452b7dc43a0f799cbaf88ff77e4e9d3a4ef`。
- 恢复服务器训练：2026-07-26 审计时，RTX 4090 24GB 上 public MLM run 已到 epoch 6，eval loss 8.6741，GPU 接近满载。
- 编辑稿件记录当前审计包包含 70 manifest、240 readable matrix files、4,544,570 referenced cells。
- 本地已有 `SnowLotus-CellFM_editor-v0.3_submit-now.zip`、中文功能创新说明、稿件草稿、cover note、README 和模型校验信息。

## 建议摘要

我们已经完成 SnowLotus-CellFM 的第一版可复现研究发布。该项目面向植物单细胞和单核转录组数据，建立了从公开数据发现、矩阵级审计、语料构建、Transformer masked modelling、层级细胞类型注释、模型冻结、外部 benchmark 准备到 GitHub/编辑包交付的完整链路。当前版本冻结了可供审阅的 annotation 与 embedding checkpoint，保留 SHA256 校验和模型卡，并将天山雪莲明确作为目标物种迁移场景处理。现阶段不夸大为已经完成雪莲单细胞图谱，而是提供一个可立即审阅、可继续训练、可在获得雪莲单细胞矩阵后快速适配的植物单细胞基础模型框架。

## 边界

- 当前审计未确认可复用的公开 Snow Lotus scRNA/snRNA cell-by-gene matrix。
- 2026-07-27 当前 Matpool 端口返回 Connection refused，服务器端口需恢复或更换；这不影响本地提交包和 GitHub 已完成内容。
- scPlantAnnotate 需要授权或结果导出，当前不能写成完成指标。
- GitHub 给编辑/审稿人使用前，需要添加访问权限或切换为 public。
"""
    MD_OUT.write_text(text, encoding="utf-8")


def main() -> None:
    build_docx()
    build_md()
    print(DOCX_OUT)
    print(MD_OUT)


if __name__ == "__main__":
    main()
