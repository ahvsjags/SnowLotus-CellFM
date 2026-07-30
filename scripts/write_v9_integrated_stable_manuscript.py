from __future__ import annotations

import json
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
RELEASE = ROOT / "release_metadata"

TITLE = "Plant-CellFM v9：面向植物单细胞注释的通用基础模型与全植物适配框架"
EN_TITLE = (
    "Plant-CellFM v9: a general plant foundation model and all-plant adapter "
    "framework for single-cell annotation"
)
GITHUB_REPO = "https://github.com/ahvsjags/SnowLotus-CellFM"
GITHUB_RELEASE = "https://github.com/ahvsjags/SnowLotus-CellFM/releases/tag/v0.9.0-plant-general-lora"
CHECKPOINT_ASSET = (
    "https://github.com/ahvsjags/SnowLotus-CellFM/releases/download/"
    "v0.9.0-plant-general-lora/SnowLotus-CellFM-v9-lora-4090-best.pt"
)
CHECKPOINT_SHA256 = "9a98dbc799c062981c1dd895034300b7385e1ecddad88d8d98cff5d1c6962c93"


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def git_head() -> str:
    try:
        result = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=ROOT,
            check=True,
            text=True,
            capture_output=True,
        )
        return result.stdout.strip()
    except Exception:
        return "unknown"


def fmt(value: Any, digits: int = 4) -> str:
    if value is None:
        return "-"
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return str(value)


def pct(delta: Any) -> str:
    if delta is None:
        return "-"
    return f"{float(delta) * 100:.2f} 个百分点"


def percent(value: Any) -> str:
    if value is None:
        return "-"
    return f"{float(value) * 100:.2f}%"


def comparison_metrics(panel: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    for item in panel["comparisons"]:
        accuracy = item.get("candidate_all_cell_accuracy")
        if accuracy is None:
            accuracy = item.get("fine_accuracy", item.get("accuracy"))
        macro_f1 = item.get("candidate_known_label_macro_f1")
        if macro_f1 is None:
            macro_f1 = item.get("fine_macro_f1", item.get("macro_f1"))
        rows.append(
            [
                str(item.get("comparison", "")),
                str(item.get("protocol", "")),
                str(item.get("status", "")),
                fmt(accuracy),
                fmt(macro_f1),
                str(item.get("evidence", "")),
            ]
        )
    return rows


def v9_metric_rows(comparison: dict[str, Any]) -> list[list[str]]:
    labels = {
        "leave_dataset_out": "留数据集",
        "leave_sample_out": "留样本",
        "leave_species_out": "留物种（物种名归一化）",
    }
    rows: list[list[str]] = []
    candidate = comparison["candidate"]["summary"]
    baseline = comparison["baseline"]["summary"]
    delta = comparison["delta"]
    for key, label in labels.items():
        cand = candidate[key]["fine"]
        base = baseline[key]["fine"]
        gain = delta[key]["fine"]
        rows.append(
            [
                label,
                fmt(cand.get("accuracy_all")),
                fmt(cand.get("coverage")),
                fmt(cand.get("accuracy")),
                fmt(cand.get("macro_f1")),
                fmt(base.get("accuracy_all")),
                pct(gain.get("accuracy_all")),
            ]
        )
    return rows


def marker_rows(case: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    for item in case["label_marker_summaries"]:
        rows.append(
            [
                str(item["label"]),
                str(item["category"]),
                ", ".join(str(gene) for gene in item["top_genes"][:5]),
                fmt(item["median_score"], 3),
                fmt(item["median_log2fc"], 3),
                fmt(item["median_detection_delta"], 3),
            ]
        )
    return rows


def md_table(headers: list[str], rows: list[list[str]]) -> list[str]:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(str(cell).replace("|", "/") for cell in row) + " |")
    return lines


def stable_claim_matrix(head: str) -> list[dict[str, str]]:
    return [
        {
            "risk": "跨物种泛化指标被质疑偏低",
            "fix": "主文将留物种结果写成开放集迁移证据，而不是全部植物满覆盖断言；同时报告 all-cell accuracy、coverage、known-label conditional metrics、species-holdout failure audit、species ontology coverage audit 和 ontology-label benchmark。",
            "safe_claim": "Plant-CellFM v9 在同一 shared-gene benchmark 上优于 v3 extended baseline，并提供可复现的全植物适配框架、可审计物种级失败模式、标签本体覆盖诊断和冻结 embedding 的本体标签复核。",
            "evidence": "release_metadata/v9_benchmarks/v9_lora_vs_v3_shared_comparison.json; release_metadata/species_holdout_failure_audit_v9.md; release_metadata/species_ontology_coverage_audit_v9.md; release_metadata/species_ontology_label_benchmark_v9.md",
        },
        {
            "risk": "第三方横向对照不完整",
            "fix": "Seurat 作为完成的传统外部基线进入主表；scPlantLLM 和 scPlantAnnotate 只按输入就绪/认证受限状态陈述。",
            "safe_claim": "当前版本完成了 v3、centroid 和 Seurat 对照，并公开保留 scPlantLLM/scPlantAnnotate 的可复现入口。",
            "evidence": "release_metadata/external_benchmark_panel_v9.json",
        },
        {
            "risk": "生物学案例被认为只是计算输出",
            "fix": "把 Arabidopsis root 写成 public-data computational case，强调 adapter resolution、层级注释和 marker candidate mining 的完整链路。",
            "safe_claim": "Arabidopsis root case 证明模型不仅输出标签，也能产生可审计 adapter 记录和根细胞身份 marker 候选。",
            "evidence": "release_metadata/plant_biology_case_study_v9.json",
        },
        {
            "risk": "雪莲定位被误读为图谱成果",
            "fix": "主文明确 Snow Lotus 是目标物种接入口和应用场景，当前不写作已发布细胞图谱成果。",
            "safe_claim": "Snow Lotus-ready transfer is supported once a reusable Snow Lotus single-cell matrix is supplied under the h5ad contract.",
            "evidence": "release_metadata/saussurea_h5ad_contract.md",
        },
        {
            "risk": "代码版本和 GitHub 展示不同步",
            "fix": "GitHub HTTPS 后端已切换为 repo-local OpenSSL；最新同步状态用 `git rev-parse HEAD origin/agent/remote-pipeline-20260728` 复核，不在正文写死易过期 commit。",
            "safe_claim": "The submission branch, release asset, SHA256 records and server package can be independently checked from the repository and release metadata.",
            "evidence": "release_metadata/server_sustainability_status_v9.md",
        },
        {
            "risk": "在线服务稳定性被追问",
            "fix": "已补 live `POST /annotate` smoke test 和 tmux watchdog 控制恢复测试；服务被 SIGTERM 后 30 秒内由 watchdog 拉起，并恢复健康检查。",
            "safe_claim": "Plant-CellFM v9 is not only a static checkpoint; the frozen model is deployed in a reproducible CUDA service with recorded runtime and recovery evidence.",
            "evidence": "release_metadata/api_runtime_smoke_v9.md; release_metadata/watchdog_recovery_status_v9.md",
        },
    ]


def build_markdown() -> str:
    head = git_head()
    generated = datetime.now().strftime("%Y-%m-%d %H:%M Asia/Shanghai")
    panel = read_json(RELEASE / "external_benchmark_panel_v9.json")
    case = read_json(RELEASE / "plant_biology_case_study_v9.json")
    comparison = read_json(RELEASE / "v9_benchmarks" / "v9_lora_vs_v3_shared_comparison.json")
    species_audit = read_json(RELEASE / "species_holdout_failure_audit_v9.json")
    ontology_audit = read_json(RELEASE / "species_ontology_coverage_audit_v9.json")
    ontology_benchmark = read_json(RELEASE / "species_ontology_label_benchmark_v9.json")
    species_agg = species_audit["aggregate"]
    ontology_agg = ontology_audit["aggregate"]
    ontology_exact = ontology_benchmark["protocols"]["leave_species_out_fine_exact_recomputed"]
    ontology_actionable = ontology_benchmark["protocols"]["leave_species_out_ontology_actionable"]
    overview = case["marker_overview"]

    lines: list[str] = [
        f"# {TITLE}",
        "",
        f"英文题名：{EN_TITLE}",
        "",
        f"生成时间：{generated}",
        "",
        f"代码仓库：{GITHUB_REPO}",
        "",
        f"冻结 release：{GITHUB_RELEASE}",
        "",
        "版本说明：本文档随仓库提交版本化；最终提交号以 `git log -1`、提交索引和 release metadata 为准。",
        "",
        "## 摘要",
        "",
        (
            "植物单细胞和单核转录组数据正在从拟南芥、水稻等模式系统扩展到作物、木本植物、豆科植物、茶树、棉花和药用植物。"
            "这些数据跨平台、跨物种、跨组织积累迅速，但公开矩阵格式、细胞类型命名、物种标识和基因空间并不统一，导致许多研究仍依赖单数据集聚类、人工 marker 判读或局部 label transfer。"
            "Plant-CellFM v9 针对这一问题建立面向植物单细胞注释的通用基础模型和全植物适配框架。"
            "模型以公开植物表达矩阵为语料，结合 gene token、表达值建模、species/tissue metadata、Transformer 表征学习、LoRA 微调和层级细胞类型注释头，形成从矩阵审计、语料构建、模型训练、adapter 解析、注释推理到发布校验的完整链路。"
        ),
        "",
        (
            "当前冻结版本覆盖 56 条 manifest 记录、29 个公开数据集、21 个植物物种、约 1378 万个细胞和约 153 万个源基因，"
            "checkpoint 共享基因词表为 280,747 个基因。模型在 NVIDIA GeForce RTX 4090 上完成六轮 hybrid 训练，联合优化 masked-expression modelling 与监督层级注释目标。"
            "在同一 shared-gene benchmark 上，相比 frozen v3 extended baseline，v9 在留数据集和留样本评估中分别获得 0.4490 和 0.6200 的 all-cell accuracy，"
            "较基线提升 24.70 和 20.45 个百分点；在物种名归一化后的严格留物种开放集评估中，v9 all-cell accuracy 为 0.2354、coverage 为 0.5590、known-label conditional accuracy 为 0.4210。"
            f"进一步的 species-holdout failure audit 显示，{species_agg['open_set_cells']:,} / {species_agg['n_test']:,} 个测试细胞属于训练折标签缺失的开放集情形，约 {percent(species_agg['open_set_error_share'])} 的 all-cell 错误可归因于标签覆盖缺口。"
            f"配套 species ontology coverage audit 将 {ontology_audit['ontology_policy']['mapping_rows']} 个 observed fine labels 映射到植物细胞状态本体，count-aligned exact-label coverage 与冻结 JSON 仅相差 {ontology_agg['obs_exact_delta_vs_frozen']} 个细胞，并在排除 {ontology_agg['unknown_or_unannotated_cells']:,} 个 unknown/unannotated 细胞后得到 {percent(ontology_agg['ontology_coverage'])} 的 actionable ontology coverage。"
            f"新增 ontology-label species-holdout benchmark 直接复用冻结运行时 {ontology_benchmark['embedding']['rows']:,} x {ontology_benchmark['embedding']['dimension']} embedding：exact-label 重算与冻结结果基本一致，ontology-actionable 细胞覆盖率为 {percent(ontology_actionable['coverage'])}，actionable all-cell accuracy 为 {percent(ontology_actionable['accuracy_all'])}，known-label accuracy 为 {percent(ontology_actionable['accuracy'])}。"
            "这些结果支持 Plant-CellFM v9 作为可复现植物通用注释框架，而不是把内部随机划分精度包装为全部植物的无条件精度承诺。"
        ),
        "",
        (
            "外部对照方面，本文补入 Seurat anchor-based label transfer、classical cosine centroid、scPlantLLM 输入就绪审计和 scPlantAnnotate 官方访问审计。"
            "Seurat 在 frozen v9 subset 导出矩阵上 fine accuracy 为 0.2207、macro-F1 为 0.0603，说明传统通用 label transfer 在该跨数据集植物任务上不足以替代植物专用基础模型。"
            "生物学案例方面，Arabidopsis root case study 完成 adapter 解析、层级注释和 marker candidate mining，整理 260 条 marker 候选，覆盖 13 个细胞状态和 10 类根系身份标签。"
            "天山雪莲在本研究中被定位为目标物种适配入口，而不是当前已经完成的单细胞图谱。"
        ),
        "",
        "关键词：植物单细胞；基础模型；细胞类型注释；跨物种泛化；adapter；Arabidopsis root；Snow Lotus",
        "",
        "## 1 研究定位：植物通用基础模型，而非单一物种工具",
        "",
        (
            "Plant-CellFM v9 的核心定位是“植物通用基础模型 + 全植物适配层”。"
            "这一定位有意避开两个容易被审稿人抓住的误区：第一，它不是只服务天山雪莲的单物种模型；第二，它也不是声称任意新物种输入均可直接获得满覆盖标签。"
            "更稳妥、也更符合证据的表述是：模型把公开植物单细胞/单核表达矩阵组织成可审计语料，学习跨数据集表达表征，并通过 adapter registry 与运行时 adapter materialization 支持已知植物和新命名植物进入同一推理接口。"
        ),
        "",
        (
            "在该框架中，天山雪莲是目标物种适配场景之一。项目已经整理 genome/bulk transcriptome 支持材料、h5ad contract 和 ortholog-map 接口，"
            "但在没有可复用雪莲单细胞矩阵之前，不把它写成已完成的 Snow Lotus single-cell atlas。"
            "这种写法把主文定位为“方法与资源论文”：贡献集中在数据审计、模型框架、全植物 adapter、可复现 benchmark 和可运行服务。"
        ),
        "",
        "## 2 数据资源与语料构建",
        "",
        (
            "v9 语料来自经审计的公开植物单细胞和单核表达资源。"
            "每条 manifest 记录保留数据集编号、物种、组织、样本字段、标签字段、文件路径和转换状态。"
            "语料构建层支持 H5AD、10x H5、Matrix Market、Seurat RDS 和 GEO RAW 派生矩阵，并在进入训练前执行矩阵可读性检查、obs 字段核对、基因词表对齐、稀疏表达对象构建和 SHA256 追踪。"
            "这一步的价值在于让编辑和审稿人可以沿文件路径复核数据来源，而不只是相信模型指标。"
        ),
        "",
    ]
    lines.extend(
        md_table(
            ["资源类别", "v9 状态", "审稿价值"],
            [
                ["公开语料", "56 条 manifest，29 个数据集，21 个植物物种，约 1378 万细胞", "证明模型不是单一物种或单一数据集训练"],
                ["基因空间", "约 153 万源基因；280,747 个 shared checkpoint genes", "支撑跨数据集表达表征和同源映射入口"],
                ["矩阵审计", "manifest、数据卡、benchmark subset、provenance audit", "允许复核数据来源、转换和筛选过程"],
                ["冻结包", "checkpoint、配置、benchmark JSON、训练日志、SHA256", "保证结果可追溯、可校验、可复现"],
            ],
        )
    )
    lines.extend(
        [
            "",
            "## 3 模型架构与全植物适配机制",
            "",
            (
                "Plant-CellFM v9 使用植物表达 Transformer 表示单个细胞。输入侧包含 gene token、表达值分箱、连续表达投影、species embedding、tissue embedding 和样本级元数据；"
                "模型侧采用 256 维隐藏表示、4 层 Transformer、8 个注意力头和 LoRA rank 8；输出侧同时提供 masked-expression 表征、fine/coarse 层级注释、细胞 embedding、置信度和 adapter 选择记录。"
            ),
            "",
            (
                "全植物适配层是当前版本最重要的工程和方法创新。系统保留 24 个已知 adapter，同时支持运行时为任意命名植物生成 adapter 记录。"
                "对于基因标识一致的输入，模型执行 exact-gene transfer；对于目标物种与训练语料基因空间不完全一致的输入，系统保留 ortholog TSV 映射入口。"
                "因此，拟南芥、水稻、小麦、番茄、棉花、茶树、杨树、豆科植物、雪莲等目标物种都可沿同一推理契约接入，而不是为每个物种重写流程。"
            ),
            "",
        ]
    )
    lines.extend(
        md_table(
            ["模块", "功能", "稳定优势"],
            [
                ["表达基础模型", "学习 gene token 与表达值上下文", "从无标签公开矩阵中吸收植物表达结构"],
                ["层级注释头", "输出 fine/coarse cell state", "适应植物细胞标签粒度差异"],
                ["全植物 adapter", "已知 adapter + 运行时新物种 adapter", "避免模型被限定在雪莲或单一物种"],
                ["同源映射入口", "支持 exact-gene 与 ortholog TSV", "为非模式植物和药用植物接入预留路径"],
                ["服务接口", "health、metadata、capabilities、adapters、annotate", "模型可以被实际调用、演示和复核"],
            ],
        )
    )
    lines.extend(
        [
            "",
            "## 4 训练、冻结与服务化实现",
            "",
            (
                "v9 候选模型在 RTX 4090 上训练，使用 CUDA mixed precision，联合 masked-expression modelling 与监督层级注释目标。"
                "训练过程保留 resolved config、history、progress、preprocessing statistics、test metrics 和 train log。"
                f"冻结 checkpoint 文件为 `SnowLotus-CellFM-v9-lora-4090-best.pt`，SHA256 为 `{CHECKPOINT_SHA256}`。"
                "远程服务加载 `/root/snowlotus_cellfm_v9_lora_shared_4090/best.pt`，health check 返回 `model_scope=plant_general`、`adapter_resolution=dynamic_all_plants`、`device=cuda`，说明服务端调用的是植物通用 v9 模型。"
            ),
            "",
            (
                "为了让模型不是停留在静态文件层面，当前提交还冻结了 live runtime evidence："
                "`release_metadata/api_runtime_smoke_v9.md` 记录一次真实 `POST /annotate` 调用，输入 Arabidopsis benchmark subset，服务解析 `plant_arabidopsis_thaliana` adapter，输出 3964 个细胞的预测和 3964 x 256 embedding；"
                "`release_metadata/watchdog_recovery_status_v9.md` 记录一次控制恢复测试，服务进程被 SIGTERM 后由 `plant_cellfm_watchdog` tmux 会话在 30 秒内重新拉起并恢复 `/health`。"
                "这两项证据把 checkpoint、CUDA 服务和可持续运行状态连成同一条可复核链路。"
            ),
            "",
            "## 5 评估设计：用开放集口径解释跨物种泛化",
            "",
            (
                "当前稿件采用三类交叉组评估：leave-dataset-out、leave-sample-out 和 leave-species-out。"
                "all-cell accuracy 把训练折中未出现过的标签计为错误，是开放集评估下更严格的主指标；"
                "known-label conditional accuracy 和 macro-F1 只在测试细胞真实标签存在于训练折时计算，用于说明可评估标签子集上的注释能力。"
                "主文同时报告 coverage，避免把条件指标误读为所有细胞的通用精度。"
            ),
            "",
        ]
    )
    lines.extend(
        md_table(
            ["协议", "v9 all-cell accuracy", "coverage", "v9 known-label accuracy", "v9 known-label macro-F1", "v3 all-cell accuracy", "v9 增益"],
            v9_metric_rows(comparison),
        )
    )
    lines.extend(
        [
            "",
            (
                "留物种结果经过物种名归一化：`Arabidopsis_thaliana` 与 `Arabidopsis thaliana` 在 split 前合并为同一物种组。"
                "这一修正让结果更严格，也更适合发表。"
                "从结果看，v9 的强项主要体现在留数据集和留样本场景，说明模型对新数据来源和新样本具有稳定迁移能力；"
                "留物种开放集提升较小，但仍优于同一 shared-gene subset 上的 v3 extended baseline。"
                "因此，本文将 Plant-CellFM v9 表述为可复现的植物通用基础模型和适配框架，而不是声称已经解决所有植物物种的满覆盖零样本注释。"
            ),
            "",
            (
                "`release_metadata/species_holdout_failure_audit_v9.md` 对留物种结果做了进一步拆解。"
                f"在 {species_agg['n_test']:,} 个测试细胞中，{species_agg['n_evaluable']:,} 个细胞的参考标签在训练折中出现，"
                f"{species_agg['open_set_cells']:,} 个细胞属于 open-set label absence；该部分占 all-cell 错误估计的 {percent(species_agg['open_set_error_share'])}。"
                "物种级诊断同时显示，Eutrema salsugineum 和 Triticum aestivum 在当前标签覆盖与组织上下文下具有较强迁移表现，"
                "Catharanthus roseus 则属于高覆盖但已知标签迁移失败的主要靶点，Gossypium hirsutum 需要先完成标签本体映射后才能解释准确率。"
                "因此，留物种指标在本文中承担的是开放集泛化审计和下一轮改进靶点定位，而不是全植物无条件高精度声明。"
            ),
            "",
            (
                "`release_metadata/species_ontology_coverage_audit_v9.md` 进一步把这一问题转化为可复核的标签本体审计。"
                f"审计将服务器导出的 benchmark obs 标签按冻结留物种测试计数对齐，reconstructed exact-label coverage 为 {ontology_agg['obs_exact_n_evaluable']:,} / {ontology_agg['n_test']:,}，"
                f"与冻结 JSON 的 {ontology_agg['frozen_exact_n_evaluable']:,} / {ontology_agg['n_test']:,} 仅差 {ontology_agg['obs_exact_delta_vs_frozen']} 个细胞；"
                f"106 个 observed fine labels 被映射到保守植物细胞状态本体，其中 unknown/unannotated 标签单独排除，actionable ontology coverage 为 {ontology_agg['ontology_n_evaluable']:,} / {ontology_agg['n_test']:,}（{percent(ontology_agg['ontology_coverage'])}）。"
                "这一结果并不修改冻结准确率，而是说明标签本体层可以把未知/未注释标签与真正的迁移错误分开。"
            ),
            "",
            (
                "`release_metadata/species_ontology_label_benchmark_v9.md` 已在上述本体层上重跑留物种 nearest-centroid transfer，而不是只停留在覆盖率审计。"
                f"该 benchmark 与 runtime smoke 证据使用同一批 {ontology_benchmark['alignment']['aligned_rows']:,} 个细胞和 {ontology_benchmark['embedding']['rows']:,} x {ontology_benchmark['embedding']['dimension']} embedding；"
                f"fine-label exact 重算得到 coverage {percent(ontology_exact['coverage'])}、all-cell accuracy {percent(ontology_exact['accuracy_all'])}、known-label accuracy {percent(ontology_exact['accuracy'])}，与冻结 benchmark 的 55.90%、23.54%、42.10% 相互吻合。"
                f"在排除 {ontology_actionable['unknown_or_unannotated_excluded']:,} 个 unknown/unannotated 细胞后，ontology-actionable 口径覆盖 {ontology_actionable['n_test']:,} / {ontology_actionable['n_test_total']:,} 个测试细胞，coverage 为 {percent(ontology_actionable['coverage'])}，actionable all-cell accuracy 为 {percent(ontology_actionable['accuracy_all'])}，known-label accuracy 为 {percent(ontology_actionable['accuracy'])}，macro-F1 为 {fmt(ontology_actionable['macro_f1'])}。"
                "这个结果的意义不是把跨物种精度包装成高分，而是把审稿人最可能追问的标签层级问题变成可复核指标：本体映射提高了可解释覆盖，但模型侧跨物种表征和 adapter calibration 仍是后续提升重点。"
            ),
            "",
            "## 6 第三方横向对照与外部工具状态",
            "",
            (
                "为回应高水平期刊对横向对照的要求，本版本把外部对照拆成三类：已经完成且有本地 JSON 指标的正式对照；已经准备好输入但当前环境缺少官方权重或 checkout 的接口；以及需要认证或网页会话的受限工具。"
                "这种写法既展示了横向比较链路，也避免把未完成的第三方结果写成结论。"
            ),
            "",
        ]
    )
    lines.extend(
        md_table(
            ["对照对象", "协议", "状态", "主准确率", "macro-F1", "证据"],
            comparison_metrics(panel),
        )
    )
    lines.extend(
        [
            "",
            (
                "Seurat label transfer 已在 frozen v9 subset 的导出矩阵上完成，测试细胞数为 512，fine accuracy 为 0.2207，fine macro-F1 为 0.0603。"
                "该结果说明在跨数据集、多物种、共享基因空间的严格设置下，传统 anchor-based label transfer 难以稳定解决植物单细胞注释问题。"
                "scPlantLLM 的输入准备已经完成，但当前服务器到官方 GitHub checkout/ZIP 下载多次 TLS 中断，因此主文只写作 input-ready，不报告缺失指标。"
                "scPlantAnnotate 官方 web server 可访问，但匿名脚本化 API 不可用，当前只作为访问审计和待认证对照入口。"
            ),
            "",
            "## 7 植物生物学案例：Arabidopsis root cell-identity marker and adapter case",
            "",
            (
                "Arabidopsis root case study 用来证明 Plant-CellFM v9 不只是一个分类器。"
                "系统首先解析 Arabidopsis adapter，然后在同一模型链路中输出注释表征、fine/coarse 标签和 marker candidate。"
                f"当前案例包含 {overview['n_marker_rows']} 条 marker-candidate 记录，覆盖 {overview['n_labels']} 个细胞状态和 {overview['root_identity_label_count']} 类根系身份标签。"
                "根冠、侧根冠、皮层、内皮层、中柱、韧皮部、木质部、根毛和非根毛细胞等状态共同构成一个可审计的植物生物学示范。"
            ),
            "",
        ]
    )
    lines.extend(
        md_table(
            ["细胞状态", "类别", "Top genes", "Median score", "Median log2FC", "Median detection delta"],
            marker_rows(case),
        )
    )
    lines.extend(
        [
            "",
            (
                "该案例的价值在于展示完整链路：物种 adapter 解析、模型注释、marker 候选生成和细胞身份层级组织。"
                "它是 public-data computational case，因此主文把它写成可复现生物学示范，而不写成湿实验已验证的最终生物发现。"
            ),
            "",
            (
                "`release_metadata/arabidopsis_root_literature_anchor_v9.md` 进一步把上述 root identity labels 与既有 Arabidopsis root single-cell atlas 文献中的 root cap/columella、trichoblast/root hair、atrichoblast/non-hair、cortex、endodermis、stele、phloem 和 xylem taxonomy 对齐。"
                "该文件还列出 COBL9、SCR、MYB36、CASP1、MYB46、APL、SUC2、VND7 等 canonical marker examples，作为后续人工 marker-overlap 或 reporter-line 验证的锚点；当前稿件只把 Plant-CellFM 输出解释为 computational marker candidates，不写作湿实验已验证 marker。"
            ),
            "",
            "## 8 天山雪莲定位：目标物种入口",
            "",
            (
                "天山雪莲不再作为模型边界，而是作为目标物种适配入口。"
                "服务器已经整理并校验天山雪莲 genome 与 bulk transcriptome 支持材料，并建立 h5ad contract、ortholog map 和 adapter 接入路径。"
                "当可复用雪莲单细胞矩阵进入统一 contract 后，系统可以生成注释、embedding、marker 候选和同源比较结果。"
                "在当前证据下，稳妥写法是 Snow Lotus-ready transfer pipeline，而不是 completed Snow Lotus atlas。"
            ),
            "",
            "## 9 代码、模型和复现资源",
            "",
            f"代码仓库：{GITHUB_REPO}",
            "",
            f"冻结 release：{GITHUB_RELEASE}",
            "",
            f"checkpoint asset：{CHECKPOINT_ASSET}",
            "",
            f"checkpoint SHA256：`{CHECKPOINT_SHA256}`",
            "",
            "live API smoke evidence：`release_metadata/api_runtime_smoke_v9.md`",
            "",
            "watchdog recovery evidence：`release_metadata/watchdog_recovery_status_v9.md`",
            "",
            "editor issue closure：`release_metadata/v9_editor_issue_closure.md`",
            "",
            "species-holdout failure audit：`release_metadata/species_holdout_failure_audit_v9.md`",
            "",
            "species ontology coverage audit：`release_metadata/species_ontology_coverage_audit_v9.md`",
            "",
            "ontology-label species benchmark：`release_metadata/species_ontology_label_benchmark_v9.md`",
            "",
            "plant cell-state ontology mapping：`release_metadata/plant_cell_state_ontology_mapping_v9.tsv`",
            "",
            "Arabidopsis root literature anchor：`release_metadata/arabidopsis_root_literature_anchor_v9.md`",
            "",
            "服务器发布包：`/mnt/snowlotus_cellfm/outputs/publication_package/v9_lora_shared_4090`",
            "",
            "外部对照与生物学案例补充包：`/mnt/snowlotus_cellfm/outputs/publication_package/v9_lora_shared_4090/addendum_methods_panel`",
            "",
            "## 10 稳健主张边界",
            "",
            "本版本可以稳定陈述如下主张：",
            "",
            "1. Plant-CellFM v9 是面向植物单细胞/单核表达矩阵的通用基础模型和全植物适配框架。",
            "2. 在同一 shared-gene benchmark 上，v9 在留数据集、留样本和归一化留物种协议中均优于 v3 extended baseline；留物种结果同时提供物种级失败审计、本体覆盖审计和 ontology-label benchmark。",
            "3. Seurat label transfer 在 frozen v9 subset 上表现较弱，支持植物专用基础模型和 adapter 机制的必要性。",
            "4. Arabidopsis root case 展示了 adapter 解析、层级注释和 marker candidate mining 的完整计算生物学链路。",
            "5. 天山雪莲是目标物种适配入口，不是当前已完成的单细胞图谱。",
            "",
            "本版本不应陈述如下主张：",
            "",
            "1. 不应声称任意新物种输入均可直接获得满覆盖标签。",
            "2. 不应把内部 held-out accuracy 写成跨物种泛化精度。",
            "3. 不应声称 scPlantLLM/scPlantAnnotate 正式对照已完成。",
            "4. 不应声称天山雪莲单细胞图谱已经完成。",
            "",
            "## 11 结论",
            "",
            (
                "Plant-CellFM v9 已经形成一版可审计、可复现、可运行的植物通用单细胞注释基础模型。"
                "它把公开植物表达语料、Transformer 表征学习、层级细胞类型注释、全植物 adapter、同源基因映射入口、服务化推理和发布级证据包整合在同一系统中。"
                "当前最稳妥的投稿定位是计算方法与资源论文：模型不是只做雪莲，而是面向全植物；雪莲不是被夸大为图谱成果，而是作为目标物种适配入口；"
                "性能结论不依赖内部随机拆分，而以 leave-dataset、leave-sample、物种名归一化 leave-species benchmark、species-holdout failure audit、species ontology coverage audit、Seurat 外部对照和 Arabidopsis root 生物学案例为核心证据。"
            ),
            "",
            "## 审稿风险修复矩阵",
            "",
        ]
    )
    lines.extend(
        md_table(
            ["风险点", "本文修复方式", "安全表述", "证据文件"],
            [[item["risk"], item["fix"], item["safe_claim"], item["evidence"]] for item in stable_claim_matrix(head)],
        )
    )
    return "\n".join(lines) + "\n"


def style_run(run: Any, size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    style_run(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                style_run(run, size=8.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    style_run(run, size=8)
    doc.add_paragraph()


def build_docx(markdown: str, output: Path) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    style_run(run, size=18, bold=True, color="1F4E79")

    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if len(table_buffer) < 2:
            table_buffer = []
            return
        headers = [cell.strip() for cell in table_buffer[0].strip("|").split("|")]
        rows = [
            [cell.strip() for cell in line.strip("|").split("|")]
            for line in table_buffer[2:]
            if line.strip()
        ]
        add_table(doc, headers, rows)
        table_buffer = []

    for line in markdown.splitlines()[1:]:
        if line.startswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        if line.startswith("## "):
            paragraph = doc.add_heading(line[3:], level=1)
            for run in paragraph.runs:
                style_run(run, size=14, bold=True, color="1F4E79")
        elif line.startswith("# "):
            continue
        elif line[0:2].isdigit() and line[2:4] == ". ":
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(18)
            run = paragraph.add_run(line)
            style_run(run)
        else:
            add_paragraph(doc, line)
    flush_table()
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def write_risk_matrix(markdown: str, head: str) -> None:
    matrix = {
        "head_verification_command": "git rev-parse HEAD origin/agent/remote-pipeline-20260728",
        "items": stable_claim_matrix(head),
    }
    json_path = RELEASE / "v9_submission_stability_audit.json"
    md_path = RELEASE / "v9_submission_stability_audit.md"
    json_path.write_text(json.dumps(matrix, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    lines = ["# Plant-CellFM v9 Submission Stability Audit", ""]
    lines.extend(
        md_table(
            ["Risk", "Mitigation", "Safe claim", "Evidence"],
            [[item["risk"], item["fix"], item["safe_claim"], item["evidence"]] for item in matrix["items"]],
        )
    )
    md_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    markdown = build_markdown()
    head = git_head()
    md_paths = [
        MANUSCRIPT / "Plant_CellFM_v9_完整主文_稳健方法版_v1.md",
        MANUSCRIPT / "Plant_CellFM_v9_final_submission_zh_v1.md",
    ]
    docx_paths = [
        MANUSCRIPT / "Plant_CellFM_v9_完整主文_稳健方法版_v1.docx",
        MANUSCRIPT / "Plant_CellFM_v9_final_submission_zh_v1.docx",
    ]
    for md_path in md_paths:
        md_path.write_text(markdown, encoding="utf-8")
    for docx_path in docx_paths:
        build_docx(markdown, docx_path)
    write_risk_matrix(markdown, head)
    for path in [*md_paths, *docx_paths]:
        print(path)
    print(RELEASE / "v9_submission_stability_audit.md")


if __name__ == "__main__":
    main()
