#!/usr/bin/env python3
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
DOC_TITLE = "Plant-CellFM v9：面向全植物单细胞注释的基础模型与多植物适配层"
GENERATED = datetime.now().strftime("%Y-%m-%d %H:%M Asia/Shanghai")

GITHUB_REPO = "https://github.com/ahvsjags/SnowLotus-CellFM"
GITHUB_RELEASE = "https://github.com/ahvsjags/SnowLotus-CellFM/releases/tag/v0.9.0-plant-general-lora"
LOCAL_COMMIT = "d6afd5e442a6876528f2cd3c084d7156f2c69ea5"
CHECKPOINT_ASSET = (
    "https://github.com/ahvsjags/SnowLotus-CellFM/releases/download/"
    "v0.9.0-plant-general-lora/SnowLotus-CellFM-v9-lora-4090-best.pt"
)
CHECKPOINT_SHA256 = "9a98dbc799c062981c1dd895034300b7385e1ecddad88d8d98cff5d1c6962c93"


def p(text: str) -> dict[str, str]:
    return {"type": "p", "text": text}


def h(text: str, level: int = 1) -> dict[str, object]:
    return {"type": "h", "text": text, "level": level}


def table(headers: list[str], rows: list[list[str]]) -> dict[str, object]:
    return {"type": "table", "headers": headers, "rows": rows}


SECTIONS: list[dict[str, object]] = [
    h("摘要", 1),
    p(
        "Plant-CellFM v9 是 SnowLotus-CellFM 项目的植物通用分支，目标不是只服务天山雪莲，"
        "而是为植物单细胞和单核转录组建立一个可训练、可审计、可复现的基础模型系统。"
        "模型面向公开植物表达矩阵学习跨数据集、跨组织和跨物种的表达表征，并通过多植物适配层"
        "把已知物种适配器、运行时新物种适配器、精细/粗粒度细胞类型注释、embedding 导出、"
        "同源基因映射和可追溯发布资产组织到同一条技术链路中。"
    ),
    p(
        "当前冻结版本为 v9 LoRA 候选模型，训练语料来自经审计的公开植物单细胞/单核表达资源，"
        "覆盖 56 条 manifest 记录、29 个公开数据集、21 个植物物种和约 1378 万个细胞；"
        "源基因规模约 153 万，checkpoint 共享基因词表为 280,747。模型在 NVIDIA RTX 4090 上完成"
        "六轮 hybrid 训练，联合优化 masked-expression modelling 与层级注释目标，并以 SHA256、"
        "模型卡、数据卡、benchmark JSON、训练历史和服务端健康检查形成可核查证据包。"
    ),
    p(
        "在严格交叉组评估中，Plant-CellFM v9 在同一 shared-gene benchmark 上相较 v3 extended 基线"
        "取得稳定提升：留数据集全细胞准确率从 0.2021 提高到 0.4490，留样本从 0.4155 提高到 0.6200。"
        "在物种名归一化后的留物种开放集评估中，模型将 `Arabidopsis_thaliana` 与 "
        "`Arabidopsis thaliana` 合并为同一物种组，保证物种拆分与生物学命名一致；该设置下 v9 全细胞准确率为 "
        "0.2354，覆盖率为 0.5590，已知标签条件准确率为 0.4210，均高于 v3 extended 基线。"
        "因此，本稿将 Plant-CellFM v9 定位为植物通用基础模型与多植物适配系统，而不是单一雪莲模型，"
        "也不把内部 held-out 指标包装成所有植物的通用精度。"
    ),
    h("1 研究定位：从单数据集注释器到植物通用基础模型", 1),
    p(
        "植物单细胞研究正在从模式植物根尖和叶片扩展到水稻、小麦、番茄、棉花、拟南芥、茶树、杨树、"
        "豆科植物和药用植物等多类系统。公开数据的快速积累带来了一个共同难题：不同项目的数据格式、"
        "物种命名、组织标签、细胞类型粒度和基因标识并不统一，使许多研究仍停留在手工 marker 判读、"
        "单数据集聚类或局部 label transfer。Plant-CellFM v9 针对这一问题建立植物通用表达基础模型，"
        "把公开矩阵、模型训练、跨物种适配、服务接口和复现材料组织为同一套可审计工程。"
    ),
    p(
        "这一定位有两个关键变化。第一，模型范围是全植物，不是天山雪莲专用；天山雪莲在本项目中是"
        "目标物种适配场景之一，和其他植物一样通过统一 h5ad contract、同源基因映射和 adapter "
        "接口进入系统。第二，模型价值不只是一组分类准确率，而是提供一个可持续吸收新植物矩阵、"
        "生成新物种适配器、输出 embedding 与注释结果、保留每一步证据链的基础设施。"
    ),
    h("2 数据资产与语料构建", 1),
    p(
        "v9 语料从公开植物单细胞与单核表达资源中筛选可读矩阵，使用 manifest 记录数据集编号、物种、"
        "组织、样本字段、标签字段、文件路径和转换状态。语料构建层支持 H5AD、10x H5、Matrix Market、"
        "Seurat RDS 与 GEO RAW 派生矩阵，并在进入训练前执行矩阵可读性检查、obs 字段核对、基因词表对齐、"
        "稀疏表达对象构建和 SHA256 追踪。"
    ),
    table(
        ["资产类别", "当前 v9 状态", "对投稿审查的价值"],
        [
            ["公开语料", "56 条 manifest、29 个数据集、21 个植物物种、约 1378 万细胞", "证明模型不是单一物种或单一数据集训练"],
            ["基因空间", "约 153 万源基因，280,747 个 shared checkpoint gene vocabulary", "支撑跨数据集表达表征和同源映射接口"],
            ["矩阵审计", "保留 manifest、数据卡、benchmark subset 和 provenance audit", "编辑和审稿人可以沿文件路径复核数据来源"],
            ["冻结包", "checkpoint、配置、benchmark JSON、训练日志、SHA256 全部打包", "保证结果不是口头描述，而是可重复核查资产"],
        ],
    ),
    h("3 模型架构与多植物适配层", 1),
    p(
        "Plant-CellFM v9 使用植物表达 Transformer 表示单细胞。输入侧包含 gene token、表达值分箱、"
        "连续表达投影、species embedding、tissue embedding 和样本级元数据；模型侧采用 256 维隐藏表示、"
        "4 层 Transformer、8 个注意力头和 LoRA rank 8；输出侧同时提供 masked-expression 表征、"
        "fine/coarse 层级注释、细胞 embedding、置信度和 adapter 选择记录。"
    ),
    p(
        "多植物适配层是当前版本的核心创新之一。系统不仅保存已知物种 adapter registry，还支持运行时为任意"
        "命名植物物种生成 adapter 记录。对于基因标识一致的输入，模型执行 exact-gene transfer；对于目标物种"
        "与训练语料基因空间不完全一致的输入，系统预留 ortholog TSV 映射入口。这样，拟南芥、水稻、小麦、番茄、"
        "棉花、茶树、杨树、豆科植物以及天山雪莲等目标物种都可以沿同一推理契约接入，而不是为每个物种重写流程。"
    ),
    table(
        ["模块", "功能", "优势"],
        [
            ["表达基础模型", "学习 gene token 与表达值上下文", "从无标签公开矩阵中吸收植物表达结构"],
            ["层级注释头", "输出 fine/coarse cell state", "适应植物细胞标签粒度差异"],
            ["多植物 adapter", "已知 adapter + 运行时新物种 adapter", "避免模型被限制在雪莲或单一物种"],
            ["同源映射入口", "支持 exact-gene 与 ortholog TSV", "为非模式植物和药用植物接入预留可执行路径"],
            ["服务接口", "health、metadata、capabilities、adapters、annotate", "模型可以被实际调用和演示"],
        ],
    ),
    h("4 训练与冻结版本", 1),
    p(
        "v9 候选模型在 RTX 4090 上训练，使用 CUDA mixed precision，联合 masked-expression modelling 与"
        "监督层级注释目标。训练过程保留 resolved config、history、progress、preprocessing statistics、"
        "test metrics 和 train log。冻结 checkpoint 文件名为 `SnowLotus-CellFM-v9-lora-4090-best.pt`，"
        f"SHA256 为 `{CHECKPOINT_SHA256}`。远程服务当前加载 `/root/snowlotus_cellfm_v9_lora_shared_4090/best.pt`，"
        "health check 返回 `model_scope=plant_general`、`adapter_resolution=dynamic_all_plants`、"
        "`device=cuda`，说明服务端调用的是植物通用 v9 模型。"
    ),
    h("5 严格评估结果与正确解读", 1),
    p(
        "当前稿件采用三类交叉组评估：leave-dataset-out、leave-sample-out 和 leave-species-out。"
        "其中 all-cell accuracy 把训练折中没有出现过的标签计为错误，是开放集评估下更严格的主指标；"
        "known-label conditional accuracy 和 macro-F1 只在测试细胞真实标签存在于训练折时计算，用于说明"
        "在可评估标签子集上的注释能力。"
    ),
    table(
        ["协议", "v9 全细胞准确率", "覆盖率", "v9 已知标签准确率", "v9 已知标签 macro-F1", "v3 extended 全细胞准确率", "v9 增益"],
        [
            ["Leave-dataset-out", "0.4490", "0.8017", "0.5601", "0.3485", "0.2021", "+0.2470"],
            ["Leave-sample-out", "0.6200", "0.9871", "0.6281", "0.4902", "0.4155", "+0.2045"],
            ["Leave-species-out（物种名归一化）", "0.2354", "0.5590", "0.4210", "0.1918", "0.1912", "+0.0441"],
        ],
    ),
    p(
        "归一化后的留物种结果比早期口径更严格。早期口径把 `Arabidopsis_thaliana` 和 `Arabidopsis thaliana` "
        "当作两个 held-out species，产生一个 100% 的小组结果，影响跨物种外推估计的可比性。当前脚本在 split 前"
        "统一将下划线替换为空格并折叠多余空白，selected benchmark 因此从 9 个 raw species labels 变为"
        "8 个 normalized species groups。这个修正使结果更保守，但更适合投稿和复现。"
    ),
    p(
        "从结果看，v9 的优势主要体现在留数据集和留样本场景，说明模型对新数据来源和新样本具有稳定迁移能力；"
        "留物种开放集提升较小，但仍优于同一 shared-gene subset 上的 v3 extended baseline。稿件因此应强调"
        "Plant-CellFM v9 是一个可复现的植物通用基础模型与适配系统，而不是声称已经解决所有植物物种的高精度"
        "零样本注释。"
    ),
    h("6 外部工具对照与可复现基准入口", 1),
    p(
        "项目已经整理 scPlantLLM 输入准备材料、scPlantAnnotate 访问审计和 benchmark package。由于部分外部"
        "工具需要授权会话、官方权重、可脚本化 API 或作者导出的结果路径，当前冻结稿不把这些第三方工具结果写成"
        "已完成结论，而是把 v3 extended baseline 作为已复现、同数据、同 subset、同指标的内部基准。"
        "这种写法避免因外部平台不可访问而被质疑，同时保留后续补充 scPlantLLM/scPlantAnnotate 正式对照的入口。"
    ),
    h("7 天山雪莲在模型中的位置", 1),
    p(
        "天山雪莲不是模型边界，而是目标物种适配案例。服务器已经下载并校验天山雪莲基因组与 bulk transcriptome "
        "支持材料，并建立 h5ad contract、同源映射和 adapter 接入路径。本版稿件将其表述为 Plant-CellFM v9 的"
        "目标物种接入能力：当雪莲单细胞矩阵进入统一 contract 后，系统可以生成注释、embedding、marker 候选和"
        "同源比较结果，从而把全植物基础模型迁移到高寒药用植物研究场景中。"
    ),
    h("8 代码、模型和发布地址", 1),
    p(
        f"代码仓库：{GITHUB_REPO}。冻结 v9 release：{GITHUB_RELEASE}。当前本地最新提交为 `{LOCAL_COMMIT}`，"
        "该提交包含开放集指标口径修正、物种名归一化 benchmark、模型卡与 README 更新。GitHub Release asset 为 "
        f"{CHECKPOINT_ASSET}。模型 SHA256 为 `{CHECKPOINT_SHA256}`。"
    ),
    p(
        "发布包位于服务器 `/mnt/snowlotus_cellfm/outputs/publication_package/v9_lora_shared_4090`，其中包含"
        " checkpoint、配置、benchmark subset、benchmark JSON、v9-v3 comparison、训练日志、模型卡、脚本和"
        " `release_sha256sums.txt`。更新后的发布包已重新执行 `sha256sum -c release_sha256sums.txt`，所有文件通过校验。"
    ),
    h("9 结论", 1),
    p(
        "Plant-CellFM v9 已形成一版可提交、可演示、可复现的植物通用单细胞注释基础模型。它把公开植物表达语料、"
        "Transformer 表征学习、层级细胞类型注释、多植物 adapter、同源基因映射入口、服务化推理和发布级证据包整合在"
        "同一系统中。当前最稳妥的投稿定位是计算方法与资源论文：模型不是只做雪莲，而是面向全植物；雪莲不是被夸大为"
        "已完成图谱，而是作为目标物种适配入口；性能结论不依赖内部随机拆分，而以 leave-dataset、leave-sample 和"
        "物种名归一化 leave-species benchmark 为核心证据。"
    ),
]


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def build_markdown() -> str:
    lines = [
        f"# {DOC_TITLE}",
        "",
        f"生成时间：{GENERATED}",
        "",
        f"GitHub 仓库：{GITHUB_REPO}",
        "",
        f"冻结 Release：{GITHUB_RELEASE}",
        "",
        f"当前代码版本：`{LOCAL_COMMIT}`",
        "",
    ]
    for item in SECTIONS:
        if item["type"] == "h":
            prefix = "#" * (int(item["level"]) + 1)
            lines.extend([f"{prefix} {item['text']}", ""])
        elif item["type"] == "p":
            lines.extend([str(item["text"]), ""])
        elif item["type"] == "table":
            lines.extend([md_table(item["headers"], item["rows"]), ""])
    return "\n".join(lines).rstrip() + "\n"


def style_run(run, size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    style_run(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table_obj = doc.add_table(rows=1, cols=len(headers))
    table_obj.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table_obj.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                style_run(run, size=9, bold=True)
    for row in rows:
        cells = table_obj.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    style_run(run, size=8.5)
    doc.add_paragraph()


def build_docx(path: Path) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(DOC_TITLE)
    style_run(run, size=18, bold=True, color="1F4E79")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"中文投稿说明稿；生成时间：{GENERATED}")
    style_run(run, size=10, bold=True)

    for line in [
        f"GitHub 仓库：{GITHUB_REPO}",
        f"冻结 Release：{GITHUB_RELEASE}",
        f"当前代码版本：{LOCAL_COMMIT}",
        f"v9 checkpoint SHA256：{CHECKPOINT_SHA256}",
    ]:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(line)
        style_run(run, size=9.5)

    for item in SECTIONS:
        if item["type"] == "h":
            paragraph = doc.add_heading(str(item["text"]), level=int(item["level"]))
            for run in paragraph.runs:
                style_run(run, size=15 if item["level"] == 1 else 12.5, bold=True, color="1F4E79")
        elif item["type"] == "p":
            add_paragraph(doc, str(item["text"]))
        elif item["type"] == "table":
            add_table(doc, item["headers"], item["rows"])

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    MANUSCRIPT.mkdir(parents=True, exist_ok=True)
    markdown = build_markdown()
    md_targets = [
        MANUSCRIPT / "Plant_CellFM_v9_植物通用基础模型_多植物适配层_中文稿_v1.md",
        MANUSCRIPT / "SnowLotus_CellFM_中文论文稿_模型功能优势详版_v0_12.md",
    ]
    docx_targets = [
        MANUSCRIPT / "Plant_CellFM_v9_植物通用基础模型_多植物适配层_中文稿_v1.docx",
        MANUSCRIPT / "SnowLotus_CellFM_中文论文稿_模型功能优势详版_v0_12.docx",
    ]
    for path in md_targets:
        path.write_text(markdown, encoding="utf-8")
    for path in docx_targets:
        build_docx(path)
    for path in [*md_targets, *docx_targets]:
        print(path)


if __name__ == "__main__":
    main()
