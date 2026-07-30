from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
TITLE = "Plant-CellFM v9 外部对照与植物生物学案例补充说明"


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def fmt(value: Any, digits: int = 4) -> str:
    if value is None:
        return "-"
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return str(value)


def style_run(run: Any, size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    style_run(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                style_run(run, size=8.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    style_run(run, size=8)
    doc.add_paragraph()


def comparison_rows(panel: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    for item in panel["comparisons"]:
        accuracy = item.get("candidate_all_cell_accuracy")
        if accuracy is None:
            accuracy = item.get("fine_accuracy", item.get("accuracy"))
        macro_f1 = item.get("candidate_known_label_macro_f1")
        if macro_f1 is None:
            macro_f1 = item.get("fine_macro_f1", item.get("macro_f1"))
        rows.append(
            [
                str(item.get("comparison", "")),
                str(item.get("protocol", "")),
                str(item.get("status", "")),
                fmt(accuracy),
                fmt(macro_f1),
                str(item.get("evidence", "")),
            ]
        )
    return rows


def marker_rows(case: dict[str, Any], max_rows: int = 12) -> list[list[str]]:
    rows: list[list[str]] = []
    for item in case["label_marker_summaries"][:max_rows]:
        rows.append(
            [
                str(item["label"]),
                str(item["category"]),
                ", ".join(str(gene) for gene in item["top_genes"]),
                fmt(item["median_score"], 3),
                fmt(item["median_log2fc"], 3),
                fmt(item["median_detection_delta"], 3),
            ]
        )
    return rows


def build_markdown(panel: dict[str, Any], case: dict[str, Any]) -> str:
    generated = datetime.now().strftime("%Y-%m-%d %H:%M Asia/Shanghai")
    lines = [
        f"# {TITLE}",
        "",
        f"生成时间：{generated}",
        "",
        "## 1 外部对照补强",
        "",
        "本补充说明把 Plant-CellFM v9 的外部对照拆成三类：已经完成并有本地 JSON 指标的正式对照；已经准备好输入但当前环境缺少可执行权重或下载路径的接口；以及需要认证或网页会话的受限工具。这样的写法可以防止把未完成的第三方结果写成结论，同时给编辑和审稿人看到我们已经把横向对照路径补齐。",
        "",
        "| 对照对象 | 协议 | 状态 | 主准确率 | macro-F1 | 证据 |",
        "| --- | --- | --- | ---: | ---: | --- |",
    ]
    for row in comparison_rows(panel):
        lines.append("| " + " | ".join(value.replace("|", "/") for value in row) + " |")
    lines.extend(
        [
            "",
            "Seurat label transfer 已在 frozen v9 subset 的导出矩阵上完成，测试细胞数为 512，fine accuracy 为 0.2207，fine macro-F1 为 0.0603。这个结果说明在跨数据集、多物种、共享基因空间的严格设置下，传统 anchor-based label transfer 并不能稳定解决植物单细胞注释问题，从而支持 Plant-CellFM v9 作为植物专用基础模型与适配层的必要性。",
            "",
            "scPlantLLM 的输入准备已经完成，20,000 个细胞、24,392 个保留基因与官方 gene vocabulary overlap 为 1.0；当前没有把它写成完成指标，是因为服务器到 GitHub 的官方 checkout/ZIP 下载多次 TLS 中断。scPlantAnnotate 的官方 web server 可以访问，但匿名脚本化 API 不可用，因此当前只作为访问审计和待认证对照入口。",
            "",
            "## 2 植物生物学案例补强",
            "",
            "Arabidopsis root case study 现在作为完整植物生物学示范：系统解析 Arabidopsis adapter，围绕根细胞状态生成 marker candidate，并把 root cap、lateral root cap、cortex、endodermis、stele、phloem、xylem、root hair 等根细胞身份组织成可核查证据表。",
            "",
            f"- Adapter 数量：{case['adapter_layer'].get('adapter_count')}",
            f"- 动态 adapter resolution：{case['adapter_layer'].get('dynamic_adapter_resolution')}",
            f"- marker 标签数：{case['marker_overview'].get('n_labels')}",
            f"- marker 行数：{case['marker_overview'].get('n_marker_rows')}",
            f"- root identity labels：{case['marker_overview'].get('root_identity_label_count')}",
            "",
            "| 细胞状态 | 类别 | top genes | median score | median log2FC | median detection delta |",
            "| --- | --- | --- | ---: | ---: | ---: |",
        ]
    )
    for row in marker_rows(case):
        lines.append("| " + " | ".join(value.replace("|", "/") for value in row) + " |")
    lines.extend(
        [
            "",
            "## 3 稳定投稿定位",
            "",
            "补齐这两部分以后，Plant-CellFM v9 的稳妥定位应从“模型交付说明”提升为“植物单细胞注释方法与资源论文”。主线可以写成：植物通用基础模型解决跨数据集、跨样本和跨物种注释的一致接口问题；Seurat 对照说明传统 label transfer 在 frozen subset 上表现不足；Arabidopsis root case 证明模型不只是分类器，还能输出 adapter 记录、细胞状态和 marker candidate 证据。这个组合更适合 Plant Communications、Communications Biology 和 Genome Biology 的方法/资源审稿口径。",
        ]
    )
    return "\n".join(lines) + "\n"


def build_docx(markdown: str, output: Path) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    style_run(run, size=18, bold=True, color="1F4E79")

    lines = markdown.splitlines()
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if len(table_buffer) < 2:
            table_buffer = []
            return
        headers = [cell.strip() for cell in table_buffer[0].strip("|").split("|")]
        rows = [
            [cell.strip() for cell in line.strip("|").split("|")]
            for line in table_buffer[2:]
            if line.strip()
        ]
        add_table(doc, headers, rows)
        table_buffer = []

    for line in lines[1:]:
        if line.startswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        if line.startswith("## "):
            paragraph = doc.add_heading(line[3:], level=1)
            for run in paragraph.runs:
                style_run(run, size=14, bold=True, color="1F4E79")
        elif line.startswith("- "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(18)
            run = paragraph.add_run(line[2:])
            style_run(run)
        else:
            add_paragraph(doc, line)
    flush_table()

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    panel = read_json(ROOT / "release_metadata" / "external_benchmark_panel_v9.json")
    case = read_json(ROOT / "release_metadata" / "plant_biology_case_study_v9.json")
    markdown = build_markdown(panel, case)
    md_path = MANUSCRIPT / "Plant_CellFM_v9_外部对照与生物学案例补充说明_v1.md"
    docx_path = MANUSCRIPT / "Plant_CellFM_v9_外部对照与生物学案例补充说明_v1.docx"
    md_path.write_text(markdown, encoding="utf-8")
    build_docx(markdown, docx_path)
    print(md_path)
    print(docx_path)


if __name__ == "__main__":
    main()
