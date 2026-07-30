from __future__ import annotations

import json
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "manuscript"
RELEASE = ROOT / "release_metadata"

REPO_URL = "https://github.com/ahvsjags/SnowLotus-CellFM"
BRANCH = "agent/remote-pipeline-20260728"
RELEASE_TAG = "v0.9.0-plant-general-lora"
CHECKPOINT_ASSET = (
    "https://github.com/ahvsjags/SnowLotus-CellFM/releases/download/"
    "v0.9.0-plant-general-lora/SnowLotus-CellFM-v9-lora-4090-best.pt"
)
CHECKPOINT_SHA256 = "9a98dbc799c062981c1dd895034300b7385e1ecddad88d8d98cff5d1c6962c93"


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def git_head() -> str:
    try:
        result = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=ROOT,
            check=True,
            text=True,
            capture_output=True,
        )
        return result.stdout.strip()
    except Exception:
        return "unknown"


def fmt(value: float) -> str:
    return f"{value:.4f}"


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def curve_at(curve: list[dict[str, Any]], rate: float) -> dict[str, Any]:
    for item in curve:
        if abs(float(item.get("acceptance_rate", -1)) - rate) < 1e-9:
            return item
    return {}


def build_context() -> dict[str, Any]:
    comparison = read_json(RELEASE / "v9_benchmarks" / "v9_lora_vs_v3_shared_comparison.json")
    ontology = read_json(RELEASE / "species_ontology_label_benchmark_v9.json")
    case = read_json(RELEASE / "plant_biology_case_study_v9.json")
    open_set = read_json(RELEASE / "open_set_calibration_v9.json")
    multi_case = read_json(RELEASE / "multispecies_scplantdb_case_v10.json")
    candidate = comparison["candidate"]["summary"]
    baseline = comparison["baseline"]["summary"]
    ontology_action = ontology["protocols"]["leave_species_out_ontology_actionable"]
    return {
        "head": git_head(),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M Asia/Shanghai"),
        "candidate": candidate,
        "baseline": baseline,
        "ontology_action": ontology_action,
        "case": case,
        "api_top30": curve_at(open_set["api_head_confidence"]["fine_confidence_curve"], 0.3),
        "api_top40": curve_at(open_set["api_head_confidence"]["fine_confidence_curve"], 0.4),
        "multi_case": multi_case,
    }


def cover_letter_md(ctx: dict[str, Any]) -> str:
    candidate = ctx["candidate"]
    baseline = ctx["baseline"]
    ontology_action = ctx["ontology_action"]
    marker_overview = ctx["case"]["marker_overview"]
    api_top30 = ctx["api_top30"]
    api_top40 = ctx["api_top40"]
    multi = ctx["multi_case"]
    return "\n".join(
        [
            "# Cover Letter",
            "",
            f"Generated: `{ctx['generated_at']}`",
            "",
            "Dear Editor,",
            "",
            "We submit **Plant-CellFM v9**, a plant-general foundation-model and all-plant adapter framework for single-cell and single-nucleus expression annotation. Plant single-cell studies now span multiple species, tissues and assay formats, but public reuse remains constrained by fragmented matrix formats, inconsistent cell-type names, uneven metadata and species-specific gene identifiers. Plant-CellFM v9 addresses this practical bottleneck by combining an audited public-plant expression corpus, a shared-gene Transformer representation model, hierarchical cell-state annotation, LoRA-based model freezing, runtime species-adapter resolution and a reproducible server-side release package.",
            "",
            "The current submission is framed as a computational method and resource. It is not a Snow Lotus-only model: Snow Lotus is treated as one target-species adapter entry point under the same h5ad and ortholog-map contract. The submitted evidence focuses on the plant-general annotation framework, reproducible model assets and a callable CUDA service.",
            "",
            "The release includes several reviewer-facing safeguards. First, all headline metrics are reported under strict grouped protocols rather than random cell splits. On the same shared-gene benchmark, Plant-CellFM v9 improves over the frozen v3 extended baseline in leave-dataset-out all-cell accuracy "
            f"({fmt(candidate['leave_dataset_out']['fine']['accuracy_all'])} versus {fmt(baseline['leave_dataset_out']['fine']['accuracy_all'])}) and leave-sample-out all-cell accuracy "
            f"({fmt(candidate['leave_sample_out']['fine']['accuracy_all'])} versus {fmt(baseline['leave_sample_out']['fine']['accuracy_all'])}). Under normalized leave-species-out evaluation, v9 reaches all-cell accuracy "
            f"{fmt(candidate['leave_species_out']['fine']['accuracy_all'])}, coverage {fmt(candidate['leave_species_out']['fine']['coverage'])} and known-label accuracy {fmt(candidate['leave_species_out']['fine']['accuracy'])}; these values are deliberately interpreted as open-set cross-species transfer evidence, not as universal high-accuracy annotation for every plant species.",
            "",
            "Second, the strict species-holdout result is accompanied by a failure audit, a 106-label plant cell-state ontology mapping and an ontology-label benchmark on the frozen runtime embeddings. After excluding unknown or unannotated labels, the ontology-actionable protocol covers "
            f"{ontology_action['n_test']:,} / {ontology_action['n_test_total']:,} cells ({pct(ontology_action['coverage'])}), with actionable all-cell accuracy {pct(ontology_action['accuracy_all'])}, known-label accuracy {pct(ontology_action['accuracy'])} and macro-F1 {fmt(ontology_action['macro_f1'])}. This diagnostic makes the remaining cross-species transfer problem explicit rather than hiding it behind label harmonization.",
            "",
            "Third, the open-set calibration audit adds a practical use layer for this strict benchmark. The deployed API annotation head reaches "
            f"{pct(api_top30['selective_accuracy'])} selective accuracy when automatically accepting the top 30% fine-confidence cells and {pct(api_top40['selective_accuracy'])} at the top 40% acceptance level. Lower-confidence and open-set-like cells are routed to manual review, ontology harmonization or species-adapter calibration rather than being converted directly into biological claims.",
            "",
            "Fourth, the submission includes a completed Seurat label-transfer comparator, classical centroid baselines and a v3 comparison. scPlantLLM and scPlantAnnotate are disclosed through official-source benchmark contracts with input packages, runner commands, missing artifacts and metric-closure rules. We therefore do not claim final numerical superiority over those tools until executable official metrics are frozen.",
            "",
            "Fifth, the Arabidopsis root and multi-species scPlantDB cases demonstrate biological use of the model output. The Arabidopsis case contains "
            f"{marker_overview['n_marker_rows']} marker-candidate rows across {marker_overview['n_labels']} cell states and {marker_overview['root_identity_label_count']} root-identity states, linking adapter resolution, hierarchical annotation and marker-candidate mining in a public-data plant root setting.",
            f" The multi-species scPlantDB case adds {multi['corpus']['cells']:,} cells, {multi['corpus']['species']} species, {multi['corpus']['tissues']} tissues and {multi['marker_record_count']} marker-candidate records as a second public-data biology demonstration.",
            "",
            "The release package is designed for direct inspection. The repository branch, model card, final manuscript, benchmark JSON files, server release verifier, release gate audit, watchdog recovery evidence and GitHub recovery note are included in the editor package. The frozen checkpoint is available from the GitHub release and is SHA256-pinned.",
            "",
            "We believe Plant-CellFM v9 will be useful to plant single-cell researchers who need a reproducible starting point for cross-dataset annotation, adapter-based target-species transfer and transparent benchmark auditing across heterogeneous public plant matrices.",
            "",
            "Sincerely,",
            "",
            "The Plant-CellFM / SnowLotus-CellFM authors",
            "",
        ]
    )


def availability() -> dict[str, Any]:
    return {
        "schema_version": "plant_cellfm_v9_data_code_availability_v1",
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M Asia/Shanghai"),
        "code_availability": {
            "repository": REPO_URL,
            "branch": BRANCH,
            "release_tag": RELEASE_TAG,
            "current_submission_entry": "SUBMISSION_INDEX_v9.md",
        },
        "model_availability": {
            "checkpoint_asset": CHECKPOINT_ASSET,
            "checkpoint_sha256": CHECKPOINT_SHA256,
            "model_card": "release_metadata/plant_cellfm_v9_model_card.md",
        },
        "data_availability": {
            "public_corpus_manifest_and_audits": [
                "release_metadata/v9_data_card.md",
                "release_metadata/corpus_provenance_audit.md",
                "release_metadata/data_integrity_audit.json",
                "release_metadata/v9_benchmarks/v9_lora_vs_v3_shared_comparison.json",
            ],
            "benchmark_and_case_files": [
                "release_metadata/species_holdout_failure_audit_v9.md",
                "release_metadata/species_ontology_coverage_audit_v9.md",
                "release_metadata/species_ontology_label_benchmark_v9.md",
                "release_metadata/open_set_calibration_v9.md",
                "release_metadata/third_party_benchmark_contract_v10.md",
                "release_metadata/plant_biology_case_study_v9.md",
                "release_metadata/arabidopsis_root_case_figure_v9.md",
                "release_metadata/multispecies_scplantdb_case_v10.md",
                "release_metadata/submission_scorecard_v11.md",
            ],
            "source_data_policy": "The release records public-source accessions, processed benchmark manifests and derived audit tables. Original public datasets remain available from their source repositories under their original access conditions.",
        },
        "server_reproducibility": {
            "server_root": "/mnt/snowlotus_cellfm",
            "final_editor_zip": "/mnt/snowlotus_cellfm/outputs/editor_submission_v9/Plant_CellFM_v9_editor_submission_final.zip",
            "server_verifier_command": "/root/miniconda3/envs/myconda/bin/python scripts/verify_v9_server_release.py --output-json release_metadata/server_release_verification_v9.json --output-md release_metadata/server_release_verification_v9.md",
            "release_gate_command": "/root/miniconda3/envs/myconda/bin/python scripts/write_release_gate_completion_audit_v9.py",
        },
        "claim_boundary": [
            "The release does not claim a completed Snow Lotus single-cell atlas.",
            "The release does not claim universal high-accuracy zero-shot annotation for every plant species.",
            "The release does not claim final official scPlantLLM or scPlantAnnotate numerical superiority without executable third-party benchmark closure.",
            "The release does not interpret 90+ evidence-readiness as 90+ raw cross-species accuracy.",
        ],
    }


def availability_md(data: dict[str, Any]) -> str:
    lines = [
        "# Data And Code Availability",
        "",
        f"Generated: `{data['generated_at']}`",
        "",
        "## Code Availability",
        "",
        f"Code repository: {data['code_availability']['repository']}",
        "",
        f"Submission branch: `{data['code_availability']['branch']}`",
        "",
        f"Release tag: `{data['code_availability']['release_tag']}`",
        "",
        "The reviewer-facing entry point is `SUBMISSION_INDEX_v9.md`.",
        "",
        "## Model Availability",
        "",
        f"Frozen checkpoint asset: {data['model_availability']['checkpoint_asset']}",
        "",
        f"Checkpoint SHA256: `{data['model_availability']['checkpoint_sha256']}`",
        "",
        f"Model card: `{data['model_availability']['model_card']}`",
        "",
        "## Data Availability",
        "",
        data["data_availability"]["source_data_policy"],
        "",
        "Primary release evidence files:",
        "",
    ]
    for item in data["data_availability"]["public_corpus_manifest_and_audits"]:
        lines.append(f"- `{item}`")
    lines.extend(["", "Benchmark and case evidence:", ""])
    for item in data["data_availability"]["benchmark_and_case_files"]:
        lines.append(f"- `{item}`")
    lines.extend(
        [
            "",
            "## Server Reproducibility",
            "",
            f"Server root: `{data['server_reproducibility']['server_root']}`",
            "",
            f"Final editor zip: `{data['server_reproducibility']['final_editor_zip']}`",
            "",
            "Verifier command:",
            "",
            "```bash",
            data["server_reproducibility"]["server_verifier_command"],
            "```",
            "",
            "Release gate command:",
            "",
            "```bash",
            data["server_reproducibility"]["release_gate_command"],
            "```",
            "",
            "## Claim Boundary",
            "",
        ]
    )
    for item in data["claim_boundary"]:
        lines.append(f"- {item}")
    lines.append("")
    return "\n".join(lines)


def style_run(run: Any, size: float = 10.5, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def md_to_docx(markdown: str, output: Path) -> None:
    doc = Document()
    for line in markdown.splitlines():
        if not line.strip():
            continue
        if line.startswith("# "):
            paragraph = doc.add_heading(line[2:], level=0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                style_run(run, size=16, bold=True, color="1F4E79")
        elif line.startswith("## "):
            paragraph = doc.add_heading(line[3:], level=1)
            for run in paragraph.runs:
                style_run(run, size=13, bold=True, color="1F4E79")
        elif line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(line[2:])
            style_run(run)
        elif line.startswith("```"):
            continue
        else:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(line)
            style_run(run)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    ctx = build_context()
    cover = cover_letter_md(ctx)
    data = availability()
    availability_text = availability_md(data)

    MANUSCRIPT.mkdir(parents=True, exist_ok=True)
    RELEASE.mkdir(parents=True, exist_ok=True)

    (MANUSCRIPT / "Plant_CellFM_v9_cover_letter.md").write_text(cover, encoding="utf-8")
    md_to_docx(cover, MANUSCRIPT / "Plant_CellFM_v9_cover_letter.docx")
    (RELEASE / "data_code_availability_v9.json").write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (RELEASE / "data_code_availability_v9.md").write_text(availability_text, encoding="utf-8")
    print(MANUSCRIPT / "Plant_CellFM_v9_cover_letter.md")
    print(MANUSCRIPT / "Plant_CellFM_v9_cover_letter.docx")
    print(RELEASE / "data_code_availability_v9.md")
    print(RELEASE / "data_code_availability_v9.json")


if __name__ == "__main__":
    main()
